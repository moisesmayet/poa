import datetime
import math
import os
import re
import pdfkit
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.files.storage import FileSystemStorage
from django.db.models import F, Value, Case, When, Count, Sum, IntegerField, CharField, BooleanField, Q
from django.db.models.functions import Concat, Coalesce
from django.http import JsonResponse, FileResponse
from django.shortcuts import render, redirect
from django.template.loader import get_template
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.views import View
from django.views.generic import TemplateView
from django.views.decorators.csrf import csrf_exempt
from config import settings
from notification.models import Parametro
from notification.views import SendMail, SendNotification, NotificationsUnreadList, GetURL, NotificationMarkAll, \
    NotificationMark
from uapa.models import TipoEstamento, Estamento, Colaborador, TableauEstamento, Tableau
from .models import POA, Objetivo, Linea, ObjetivoOperativo, Meta, Actividad, MedioVerificacion, Mes, \
    Responsable, Cronograma, Eje, Evidencia, Nota, Noticia, Log
from .forms import FormPOA, FormObjetivoOperativo, FormMeta, FormActividad
from authentication.models import CustomUser
from spire.doc import Document
from spire.doc import FileFormat
from spire.doc import PageOrientation


class Home(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        noticias = Noticia.objects.all().order_by('noticia_start')
        return render(request, "home.html", {"noticias": noticias})


class OrganizationChart(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        estamento_tipos = TipoEstamento.objects.all()
        estamento_tipo = TipoEstamento.objects.filter(tipo_code='REC').first()
        estamento_root = Estamento.objects.filter(estamento_tipo=estamento_tipo).first()
        return render(request, "organization_chart.html",
                      {"estamento_root": estamento_root, "estamento_tipos": estamento_tipos})


class ComponenteChart(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        ejes_list = Eje.objects.all()
        return render(request, "componente_chart.html", {"ejes_list": ejes_list})


class About(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        return render(request, "about.html")


class Help(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        return render(request, "help.html")


class ObjetivosList:
    # noinspection PyMethodMayBeStatic
    def get_objetivos_list(self, poa_root, user, estamento_id, poa_anno, poas_subs, tableros_subs, pag, filtred_by):
        poas_list = []
        tableros_list = []
        if poa_root:
            notas_poa = GetNotasPoa(poa_root.id)
            users_roots = getUsersRoots(poa_root.poa_estamento, user)
            if pag == 0 or pag == 1:  # pag == 0 es para cuando esta en edit
                poas_list = [
                    {'poa': poa_root, 'notas_poa': notas_poa, 'poa_porciento': '0%', 'poa_porciento_parcial': '0%',
                     'users_roots': users_roots}]
                tableros_list = list(
                    TableauEstamento.objects.filter(tablero_estamento=poa_root.poa_estamento,
                                                    tablero_tableau__tableau_anno=poa_root.poa_anno))

        includeSubs = IncludeSubs(estamento_id, poa_anno)
        if includeSubs:
            poas_list += poas_subs

        if filtred_by == "tablero" or includeSubs:
            tableros_list += tableros_subs

        objetivos_poa_list = []
        actividades_sorted = []
        for item in poas_list:
            poa = item['poa']
            if poa.id > 0:
                if filtred_by != "order":
                    objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id).annotate(
                        operativo_index=Concat(
                            F('operativo_order'), Value('.'),
                            output_field=CharField()
                        ))
                    objetivos_list = []
                    suma_porciento_meta = 0
                    suma_porciento_parcial_meta = 0
                    total_meta = 0
                    total_parcial_meta = 0
                    for objetivo in objetivos:
                        metas = Meta.objects.filter(meta_operativo_id=objetivo.id).annotate(
                            meta_index=Concat(
                                F('meta_operativo__operativo_order'), Value('.'),
                                F('meta_order'), Value('.'),
                                output_field=CharField()
                            ))
                        #if metas.count() > 0:
                        metas_list = []
                        for meta in metas:
                            actividades = Actividad.objects.filter(actividad_meta_id=meta.id).annotate(
                                actividad_index=Concat(
                                    F('actividad_meta__meta_operativo__operativo_order'), Value('.'),
                                    F('actividad_meta__meta_order'), Value('.'),
                                    F('actividad_order'), Value('.'),
                                    output_field=CharField()
                                ))
                            #if actividades.count() > 0:
                            actividades_list = []
                            total_meta += 1
                            porciento_meta = 0
                            porciento_parcial_meta = 0
                            peso_parcial_meta = 0
                            presupuesto = 0
                            presupuesto_gastado = 0
                            for actividad in actividades:
                                porcientos = CalcularPorcientos(actividad, presupuesto, porciento_meta,
                                                                porciento_parcial_meta, peso_parcial_meta)
                                presupuesto = porcientos[0]
                                presupuesto_gastado += porcientos[1]
                                porciento_meta = porcientos[2]
                                porciento_parcial_meta = porcientos[4]
                                peso_parcial_meta = porcientos[6]

                                meses_cronograma_list = []
                                meses_list = Mes.objects.all()
                                for mes in meses_list:
                                    cronograma = Cronograma.objects.filter(cronograma_actividad_id=actividad.id,
                                                                           cronograma_mes_id=mes.id).first()
                                    if cronograma is None:
                                        meses_cronograma_list.append('')
                                    else:
                                        meses_cronograma_list.append('X')

                                actividades_sorted.append(actividad)
                                notas = GetNotas(poa.id, actividad.id, "actividad", user)
                                actividades_list.append(
                                    {'actividad': actividad, 'meses_cronograma_list': meses_cronograma_list,
                                     'notas': notas})

                            porciento_meta = GetFormatPorciento(porciento_meta)
                            porciento_parcial_meta = GetFormatPorciento(porciento_parcial_meta)
                            suma_porciento_meta += porciento_meta
                            if peso_parcial_meta:
                                suma_porciento_parcial_meta += porciento_parcial_meta * 100 / peso_parcial_meta
                                suma_porciento_parcial_meta = round(suma_porciento_parcial_meta, 2)
                                total_parcial_meta += 1

                            notas = GetNotas(poa.id, meta.id, "meta", user)
                            metas_list.append(
                                {'meta': meta, 'porciento': porciento_meta,
                                 'porciento_parcial': porciento_parcial_meta,
                                 'presupuesto': presupuesto, 'presupuesto_gastado': presupuesto_gastado,
                                 'actividades_list': actividades_list, 'notas': notas})

                        notas = GetNotas(poa.id, objetivo.id, "objetivo", user)
                        objetivos_list.append({'objetivo': objetivo, 'metas_list': metas_list, 'notas': notas})

                    if len(objetivos_list) > 0:
                        if total_meta > 0:
                            poa_porciento = suma_porciento_meta / total_meta
                            poa_porciento = round(poa_porciento, 2)
                        else:
                            poa_porciento = 0
                        if total_parcial_meta > 0:
                            poa_porciento_parcial = suma_porciento_parcial_meta / total_parcial_meta
                            poa_porciento_parcial = round(poa_porciento_parcial, 2)
                        else:
                            poa_porciento_parcial = 0

                        poa_porciento = GetFormatPorciento(poa_porciento)
                        poa_porciento_parcial = GetFormatPorciento(poa_porciento_parcial)
                        item['poa_porciento'] = f'{poa_porciento}%'
                        item['poa_porciento_parcial'] = f'{poa_porciento_parcial}%'
                        objetivos_poa_list.append(
                            {'estamento_name': poa.poa_estamento.estamento_name, 'objetivos_list': objetivos_list,
                             'poa_porciento': f'{poa_porciento}%',
                             'poa_porciento_parcial': f'{poa_porciento_parcial}%'})
                else:
                    objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id).annotate(
                        operativo_portlet=Concat(
                            Value('portlet_objetivo_'),
                            F('id'),
                            output_field=CharField()
                        ),
                        operativo_metas=Count('meta', distinct=True) + 1,
                        operativo_index=Concat(
                            F('operativo_order'), Value('.'),
                            output_field=CharField()
                        )
                    ).order_by('operativo_order')

                    if objetivos.count() > 0:
                        metas_list = []
                        actividades_list = []
                        for objetivo in objetivos:
                            objetivo.objetivo_notas = GetNotas(poa.id, objetivo.id, "objetivo", user)
                            metas = Meta.objects.filter(meta_operativo_id=objetivo.id).order_by(
                                'meta_order'
                            ).annotate(
                                meta_portlet=Concat(
                                    Value('portlet_objetivo_'),
                                    F('meta_operativo_id'),
                                    Value('_meta_'),
                                    F('id'),
                                    output_field=CharField()
                                ),
                                meta_actividades=Count('actividad', distinct=True) + 1,
                                meta_peso=Coalesce(Sum('actividad__actividad_peso'), Value(0)),
                                meta_presupuesto=Sum('actividad__actividad_presupuesto'),
                                meta_index=Concat(
                                    F('meta_operativo__operativo_order'), Value('.'),
                                    F('meta_order'), Value('.'),
                                    output_field=CharField()
                                )
                            ).order_by('meta_order')
                            metas_list.append(
                                {'objetivo_id': objetivo.id, 'no_metas': len(metas), 'metas': list(metas)})

                            for meta in metas:
                                meta.meta_notas = GetNotas(poa.id, meta.id, "meta", user)
                                actividades = Actividad.objects.filter(actividad_meta_id=meta.id).order_by(
                                    'actividad_order'
                                ).annotate(
                                    actividad_portlet=Concat(
                                        Value('portlet_objetivo_'),
                                        F('actividad_meta__meta_operativo_id'),
                                        Value('_meta_'),
                                        F('actividad_meta_id'),
                                        Value('_actividad_'),
                                        F('id'),
                                        output_field=CharField()
                                    ),
                                    meta_peso=Value(meta.meta_peso, output_field=IntegerField()),
                                    actividad_index=Concat(
                                        F('actividad_meta__meta_operativo__operativo_order'), Value('.'),
                                        F('actividad_meta__meta_order'), Value('.'),
                                        F('actividad_order'), Value('.'),
                                        output_field=CharField()
                                    )
                                ).order_by('actividad_order')

                                for actividad in actividades:
                                    actividad.actividad_notas = GetNotas(poa.id, actividad.id, "actividad", user)
                                actividades_list.append({'meta_id': meta.id, 'actividades': list(actividades)})

                        objetivos = list(objetivos)
                        metas_list = list(metas_list)
                        actividades_list = list(actividades_list)
                        objetivos_poa_list.append(
                            {'estamento_name': poa.poa_estamento.estamento_name,
                             'objetivos': objetivos,
                             'metas_list': metas_list,
                             'actividades_list': actividades_list
                             })
                    else:
                        objetivos_poa_list.append(
                            {'estamento_name': poa.poa_estamento.estamento_name,
                             'objetivos': [],
                             'metas_list': [],
                             'actividades_list': []
                             })

        return {'objetivos_poa_list': objetivos_poa_list,
                'actividades_list': actividades_sorted,
                'poas_list': poas_list, 'tableros_list': tableros_list}

    # noinspection PyMethodMayBeStatic
    def get_ejes_list(self, poa_root, user, estamento_id, poa_anno, poas_subs, pag):
        poas_list = []
        if poa_root:
            users_roots = getUsersRoots(poa_root.poa_estamento, user)
            if pag == 1:
                poas_list = [{'poa': poa_root, 'notas_poa': '', 'poa_porciento': '0%', 'poa_porciento_parcial': '0%',
                              'users_roots': users_roots}]

        if IncludeSubs(estamento_id, poa_anno):
            poas_list += poas_subs

        operativo_index = 0
        presupuesto_total = 0
        ejes = []
        objetivos = []
        lineas = []
        poa_ejes_list = []
        for item in poas_list:
            ejes_list = []
            poa = item['poa']
            objetivos_operativos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id)
            for operativo in objetivos_operativos:
                if operativo.operativo_linea_id not in lineas:
                    lineas.append(operativo.operativo_linea_id)
                    if operativo.operativo_linea.linea_objetivo_id not in objetivos:
                        objetivos.append(operativo.operativo_linea.linea_objetivo_id)
                        if operativo.operativo_linea.linea_objetivo.objetivo_eje_id not in ejes:
                            ejes.append(operativo.operativo_linea.linea_objetivo.objetivo_eje_id)

            ejes_filtred = Eje.objects.filter(id__in=ejes)
            for eje in ejes_filtred:
                objetivos_list = []
                objetivos_filtred = Objetivo.objects.filter(id__in=objetivos, objetivo_eje=eje)
                for objetivo in objetivos_filtred:
                    high_objetivo = 0
                    lineas_list = []
                    lineas_filtred = Linea.objects.filter(id__in=lineas, linea_objetivo=objetivo)
                    for linea in lineas_filtred:
                        high_linea = 0
                        operativos_list = []
                        operativos_filtred = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id,
                                                                              operativo_linea=linea)
                        for operativo in operativos_filtred:
                            high_operativo = 0
                            metas_list = []
                            metas_filtred = Meta.objects.filter(meta_operativo=operativo)
                            for meta in metas_filtred:
                                if Actividad.objects.filter(actividad_meta=meta).count() > 0:
                                    high_objetivo += 1
                                    high_linea += 1
                                    high_operativo += 1
                                    presupuesto = 0
                                    actividades = Actividad.objects.filter(actividad_meta=meta).exclude(
                                        actividad_presupuesto=0)
                                    for actividad in actividades:
                                        presupuesto += actividad.actividad_presupuesto
                                    presupuesto_total += presupuesto
                                    metas_list.append({'meta': meta, 'presupuesto': presupuesto})
                            if len(metas_list) > 0:
                                operativo_index += 1
                                operativos_list.append(
                                    {'operativo': operativo, 'metas_list': metas_list, 'high_operativo': high_operativo,
                                     'operativo_index': operativo_index})
                        if len(operativos_list) > 0:
                            lineas_list.append(
                                {'linea': linea, 'operativos_list': operativos_list, 'high_linea': high_linea})
                    if len(lineas_list) > 0:
                        objetivos_list.append(
                            {'objetivo': objetivo, 'lineas_list': lineas_list, 'high_objetivo': high_objetivo})
                if len(objetivos_list) > 0:
                    ejes_list.append({'eje': eje, 'objetivos_list': objetivos_list})

            poa_ejes_list.append(
                {'estamento_name': poa.poa_estamento.estamento_name,
                 'ejes_list': ejes_list
                 })

        return [poa_ejes_list, presupuesto_total]


class POANotas(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        try:
            nota_id = request.POST.get('nota_id', '')
            nota_description = request.POST.get('nota_description', '')
            if nota_description or nota_id:
                nota = CrearNota(request, nota_description, True, True)
                if nota:
                    data = {
                        'success': True,
                        'nota_id': nota.id,
                        'nota_user': nota.nota_user.username.title(),
                        'nota_date': nota.nota_date.strftime('%d %b %Y %H:%M'),
                        'nota_description': nota_description,
                    }
                    return JsonResponse(data)
            return JsonResponse({'success': False, 'error': 'Error enviando la nota'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Error enviando la nota'})


class PopoverNotas(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        if request.method == 'GET':
            poa_id = request.GET['poa_id']
            nota_itemid = request.GET['nota_itemid']
            nota_itemname = request.GET['nota_itemname']
            notas = Nota.objects.filter(nota_poa_id=poa_id, nota_itemid=nota_itemid,
                                        nota_itemname=nota_itemname,
                                        nota_checked=False)
            notas_list = [GetNotaValue(nota, notas.count(), request.user, "", "") for nota in notas]
            return JsonResponse(notas_list, safe=False)
        else:
            return JsonResponse({'success': False, 'error': 'Error enviando la nota'})

    def post(self, request, *args, **kwargs):
        try:
            request_from = request.POST["request_from"]
            nota_description = request.POST["nota_description"]
            nota_id = request.POST["nota_id"]
            if nota_description or nota_id:
                updatedValue = UpdateNota(request, nota_description, nota_id, request_from)
                return JsonResponse({'updatedValue': updatedValue})
        except Exception as e:
            messages.error(request, F'Ocurrió un error. ' + str(e))

        return redirect(request_from, estamento_id=self.estamento_id, poa_anno=self.poa_anno)


class CronogramaPresupuesto(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request):
        if request.method == 'GET':
            cronograma_id = request.GET['cronograma_id']
            cronograma = Cronograma.objects.filter(id=cronograma_id).first()
            total_presupuesto = Cronograma.objects.filter(
                cronograma_actividad=cronograma.cronograma_actividad,
                cronograma_presupuesto__gt=0
            ).exclude(id=cronograma_id).aggregate(total=Sum('cronograma_presupuesto')).get('total', 0)
            if total_presupuesto is None:
                total_presupuesto = 0
            cronograma_presupuesto = cronograma.cronograma_actividad.actividad_presupuesto - total_presupuesto
            if cronograma_presupuesto > 0:
                es_negativo = False
            else:
                es_negativo = True
            cronograma_presupuesto = GetFormatPesos(cronograma_presupuesto)
            data = {
                'cronograma_presupuesto': cronograma_presupuesto,
                'es_negativo': es_negativo
            }

            return JsonResponse(data)
        else:
            return JsonResponse({'success': False, 'error': 'Error enviando la nota'})


class PoaSendMessage(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        try:
            estamento_id = request.POST.get('id', '')
            notificacion_type = request.POST.get('notificacion_type', '')
            notificacion_user = request.POST.get('notificacion_user', '')
            notificacion_message = request.POST.get('notificacion_message', '')
            notificacion_user = CustomUser.objects.get(id=notificacion_user)
            SendNotification(request.user, notificacion_user, notificacion_type, notificacion_message)
            firma = GetURL(request) + "/static/images/logo/popin_firma.png"
            SendMail(notificacion_user.username,
                     notificacion_user.email,
                     f'{str(request.user.username).title()} le envió un mensaje',
                     'basic_email.html',
                     notificacion_message,
                     '',
                     firma)
            return JsonResponse({'success': True, 'menu_estamento_id': estamento_id})
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Error enviando el mensaje'})


class RefreshNotifications(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def get(self, request, *args, **kwargs):
        try:
            notifications = NotificationsUnreadList(request.user)
            notifications_number = len(notifications)

            data = {
                'success': True,
                'notifications_number': notifications_number,
                'notifications': notifications,
            }
            return JsonResponse(data)
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Error enviando el mensaje'})


class MarkNotifications(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        try:
            notificacion_id = request.POST.get('notificacion_id', '')
            notificacion_read = request.POST.get('notificacion_readed', '')
            if notificacion_id != '' and notificacion_read != '':
                notificacion_id = int(notificacion_id)
                if notificacion_read == 'true':
                    notificacion_read = True
                else:
                    notificacion_read = False

                notifications_number = 0
                if notificacion_id == 0:
                    NotificationMarkAll(request.user, notificacion_read)
                elif notificacion_id > 0:
                    NotificationMark(notificacion_id, notificacion_read)
                    if notificacion_read:
                        notifications = NotificationsUnreadList(request.user)
                        notifications_number = len(notifications)

                data = {
                    'success': True,
                    'notifications_number': notifications_number,
                    'notification_id': notificacion_id,
                    'notification_read': notificacion_read,
                }
                return JsonResponse(data)
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Error enviando el mensaje'})


class POALoad(View):
    estamento_id = 0
    poa_anno = 0

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.estamento_id = request.POST.get('selectlevel2')
        self.poa_anno = request.POST.get('selectlevel3')
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        estamento_id = self.estamento_id
        poa_anno = self.poa_anno

        request_from = request.POST.get('request_from', 'poa_card')

        if request_from == 'poa_edit' or request_from == 'poa_preview':
            estamento = Estamento.objects.filter(id=estamento_id, estamento_user=request.user).first()
            colaborador = Colaborador.objects.filter(colaborador_estamento_id=estamento_id,
                                                     colaborador_user=request.user).first()
            poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
            if not ((estamento and estamento.estamento_has_poa)
                    or (colaborador and colaborador.colaborador_estamento.estamento_has_poa)) \
                    or (poa and poa.poa_estado_id > 2):
                if request_from == 'poa_edit':
                    request_from = 'poa_card'
            else:
                if request_from == 'poa_preview':
                    return redirect('poa_preview', estamento_id=estamento_id, poa_anno=poa_anno, pag=1)
                else:
                    return redirect('poa_edit', estamento_id=estamento_id, poa_anno=poa_anno)
        elif request_from == 'poa_cronograma':
            return redirect('poa_cronograma', estamento_id=estamento_id, poa_anno=poa_anno)

        return redirect(request_from, estamento_id=estamento_id, poa_anno=poa_anno, pag=1)


class PoaPreview(TemplateView):
    template_name = 'poa_preview.html'
    estamento_id = 0
    poa_anno = 0
    pag = 0
    send_values = None

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.estamento_id = kwargs.get('estamento_id')
        self.poa_anno = kwargs.get('poa_anno')
        self.pag = kwargs.get('pag')

        # Redirección antes de renderizar la plantilla
        self.send_values = GetPOAList(request, request.user, self.estamento_id, self.poa_anno, self.pag, "")
        if self.send_values[0] == "home":
            return redirect("home")
        elif self.send_values[0] == "poa_edit_start":
            return redirect("poa_edit_start", estamento_id=self.estamento_id, poa_anno=self.poa_anno)

        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def get_context_data(self, **kwargs):
        estamento_id = self.estamento_id
        poa_anno = self.poa_anno
        pag = self.pag
        send_values = self.send_values

        estamento = Estamento.objects.get(id=estamento_id)
        clone_annos = getCloneYears(estamento_id)

        context = super().get_context_data(**kwargs)
        context['estamento_id'] = estamento_id
        context['poa_anno'] = poa_anno
        context['estamento_name'] = estamento.estamento_name
        context['estamento_tipo_id'] = estamento.estamento_tipo_id
        context['estamentos_select'] = send_values[1]
        context['estamentos_list'] = send_values[2]
        context['estamentos_options_todos'] = send_values[3]
        context['can_edit'] = send_values[4]
        context['edit_notas'] = send_values[5]
        context['new_years'] = send_values[6]
        context['poa_years'] = send_values[7]
        context['poa'] = send_values[8]
        context['objetivos_list'] = send_values[9]
        context['notas_poa'] = send_values[12]
        context['tablero_list'] = send_values[13]
        context['poa_list'] = send_values[14]
        context['users_roots'] = send_values[15]
        context['paginas'] = send_values[16]
        context['pag'] = pag
        context['clone_annos'] = clone_annos
        context['request_from'] = "poa_preview"

        return context

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        action = str(request.POST.get('action'))
        if action == "order_all":
            try:
                request_from = request.POST.get('request_from')
                OrderAll(self.estamento_id, self.poa_anno, request.user, request_from)
            except Exception as e:
                messages.error(request, str(e))

        return redirect("poa_preview", estamento_id=self.estamento_id, poa_anno=self.poa_anno, pag=self.pag)


class PoaTablero(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        pag = kwargs.get('pag')

        poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        if not (poa and poa.poa_estamento.estamento_has_poa) and not IsRoot(estamento_id):
            return redirect("poa_card", estamento_id=estamento_id, poa_anno=poa_anno)

        send_values = GetPOAList(request, user, estamento_id, poa_anno, pag, "tablero")

        if send_values[0] == "home":
            return redirect("home")
        else:
            if send_values[0] == "poa_edit_start":
                return redirect("poa_edit_start", estamento_id=estamento_id, poa_anno=poa_anno)

        estamento = Estamento.objects.get(id=estamento_id)
        clone_annos = getCloneYears(estamento_id)
        return render(request, "poa_tablero.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "estamento_name": estamento.estamento_name,
                       "estamento_tipo_id": estamento.estamento_tipo_id, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2],
                       "estamentos_options_todos": send_values[3], "can_edit": send_values[4],
                       "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8],
                       "objetivos_list": send_values[9], "notas_poa": send_values[12], "tablero_list": send_values[13],
                       "poa_list": send_values[14], "users_roots": send_values[15], "paginas": send_values[16],
                       "pag": pag, "clone_annos": clone_annos, "request_from": "poa_tablero"})


class PoaTableau(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        tableau_id = kwargs.get('tableau_id')
        tablero = Tableau.objects.filter(id=tableau_id).first()
        return render(request, "poa_tableau.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "tablero": tablero})


class POACard(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        pag = kwargs.get('pag')
        send_values = GetPOAList(request, user, estamento_id, poa_anno, pag, "card")

        if send_values[0] == "home":
            return redirect("home")
        else:
            if send_values[0] == "poa_edit_start":
                return redirect("poa_edit_start", estamento_id=estamento_id, poa_anno=poa_anno)

        estamento = Estamento.objects.get(id=estamento_id)
        poa_list = send_values[14]
        if len(poa_list) == 0:
            poa = POA()
            poa.id = 0
            poa.poa_estado_id = 0
            poa.poa_estamento = estamento
            poa.poa_color = '#B7BBCA'
            poa.poa_user_owner = IsPOAUser(estamento_id, user)
            users_roots = getUsersRoots(poa.poa_estamento, user)
            poa_list.append({'poa': poa, 'notas_poa': '', 'poa_porciento': '0%', 'poa_porciento_parcial': '0%', 'users_roots': users_roots})
        clone_annos = getCloneYears(estamento_id)

        return render(request, "poa_card.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "estamento_name": estamento.estamento_name,
                       "estamento_tipo_id": estamento.estamento_tipo_id, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2],
                       "estamentos_options_todos": send_values[3], "can_edit": send_values[4],
                       "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8],
                       "objetivos_list": send_values[9], "notas_poa": send_values[12], "tablero_list": send_values[13],
                       "users_roots": send_values[15], "paginas": send_values[16], "pag": pag,
                       "request_from": "poa_card", "poa_list": poa_list, "clone_annos": clone_annos})


class PoaEjecutoria(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        pag = kwargs.get('pag')
        send_values = GetPOAList(request, user, estamento_id, poa_anno, pag, "")

        if send_values[0] != "":
            return redirect("home")

        estamento = Estamento.objects.get(id=estamento_id)
        clone_annos = getCloneYears(estamento_id)
        return render(request, "poa_ejecutoria.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "estamento_name": estamento.estamento_name,
                       "estamento_tipo_id": estamento.estamento_tipo_id, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2], "estamentos_options_todos": send_values[3],
                       "can_edit": send_values[4], "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8],
                       "objetivos_list": send_values[9], "notas_poa": send_values[12], "tablero_list": send_values[13],
                       "poa_list": send_values[14], "users_roots": send_values[15], "paginas": send_values[16],
                       "pag": pag, "clone_annos": clone_annos, "request_from": "poa_ejecutoria"})


class PoaCorrelacion(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        pag = kwargs.get('pag')
        send_values = GetPOAList(request, user, estamento_id, poa_anno, pag, "ejes")

        if send_values[0] != "":
            return redirect("home")

        estamento = Estamento.objects.get(id=estamento_id)
        clone_annos = getCloneYears(estamento_id)
        return render(request, "poa_correlacion.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "estamento_name": estamento.estamento_name,
                       "estamento_tipo_id": estamento.estamento_tipo_id, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2], "estamentos_options_todos": send_values[3],
                       "can_edit": send_values[4], "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8], "objetivos_list": send_values[9],
                       "notas_poa": send_values[12], "tablero_list": send_values[13], "poa_list": send_values[14],
                       "users_roots": send_values[15], "paginas": send_values[16], "pag": pag,
                       "clone_annos": clone_annos, "request_from": "poa_correlacion"})


class PoaCronogramaUpdate(View):
    cronograma_id = 0

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.cronograma_id = kwargs.get('cronograma_id')
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        cronograma_id = self.cronograma_id
        try:
            cronograma = Cronograma.objects.get(id=cronograma_id)
            if cronograma.cronograma_cumplimiento:
                cronograma.cronograma_cumplimiento = False
                cronograma.cronograma_cumplimiento_mes_id = cronograma.cronograma_mes_id
                cronograma.cronograma_cumplimiento_date = timezone.now()
                cumplimiento = '"Planificada"'
            else:
                cronograma.cronograma_cumplimiento = True
                cronograma_cumplimiento_date = cronograma.cronograma_cumplimiento_date
                date_now = timezone.now()
                cumplimiento = '"Realizada"'
                if not cronograma_cumplimiento_date or (
                        cronograma_cumplimiento_date.year != date_now.year or
                        cronograma_cumplimiento_date.month != date_now.month):
                    cronograma.cronograma_cumplimiento_date = date_now
                    cronograma.cronograma_cumplimiento_mes_id = timezone.now().month

            evidencia = Evidencia.objects.filter(evidencia_cronograma=cronograma).first()
            if evidencia and evidencia.evidencia_file != '' and os.path.exists(evidencia.evidencia_file.path):
                has_evidencia = True
            else:
                has_evidencia = False

            cronograma.save()

            cronograma_defasado = False
            if cronograma.cronograma_cumplimiento_mes_id != cronograma.cronograma_mes:
                cronograma_defasado = True

            poa = cronograma.cronograma_actividad.actividad_meta.meta_operativo.operativo_poa
            index = f'{cronograma.cronograma_actividad.actividad_meta.meta_operativo.operativo_order}.' \
                    f'{cronograma.cronograma_actividad.actividad_meta.meta_order}.' \
                    f'{cronograma.cronograma_actividad.actividad_order}.'
            RegistrarLog(poa.poa_estamento_id, poa.poa_anno, request.user, 'Update',
                         f'La actividad {index} con id {cronograma.cronograma_actividad_id} '
                         f'del mes {cronograma.cronograma_mes.mes_name} se marcó como {cumplimiento}')

            return JsonResponse({'success': True,
                                 'cronograma_cumplimiento': cronograma.cronograma_cumplimiento,
                                 'cronograma_defasado': cronograma_defasado,
                                 'has_evidencia': has_evidencia})
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Valor no encontrado'})


class PoaCronograma(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')

        poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        if not (poa and poa.poa_estamento.estamento_has_poa) and not IsRoot(estamento_id):
            return redirect("poa_card", estamento_id=estamento_id, poa_anno=poa_anno, pag=1)

        send_values = GetPOAList(request, user, estamento_id, poa_anno, 0, "actividades")

        if send_values[0] == "home":
            return redirect("home")
        else:
            if send_values[0] == "poa_edit_start":
                return redirect("poa_edit_start", estamento_id=estamento_id, poa_anno=poa_anno)

        objetivos = ObjetivoOperativo.objects.filter(operativo_poa=send_values[8])
        metas_list = Meta.objects.filter(meta_operativo__in=objetivos).annotate(
            operativo_order=F('meta_operativo__operativo_order')
        ).order_by('operativo_order', 'meta_order')

        mes_actual = Mes.objects.get(id=datetime.date.today().month)
        if datetime.date.today().month < 12:
            mes_siguiente = True
        else:
            mes_siguiente = False
        if datetime.date.today().month > 1:
            mes_anterior = True
        else:
            mes_anterior = False

        estamento = Estamento.objects.get(id=estamento_id)
        clone_annos = getCloneYears(estamento_id)
        url_base = GetURL(request)
        return render(request, "poa_cronograma.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "estamento_name": estamento.estamento_name,
                       "estamento_tipo_id": estamento.estamento_tipo_id, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2], "estamentos_options_todos": send_values[3],
                       "can_edit": send_values[4], "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8], "objetivos_list": send_values[9],
                       "cronogramas_list": send_values[11], "notas_poa": send_values[12],
                       "tablero_list": send_values[13], "poa_list": send_values[14], "users_roots": send_values[15],
                       "paginas": send_values[16], "metas_list": metas_list, "url_base": url_base,
                       "mes_actual": mes_actual, "mes_siguiente": mes_siguiente, "mes_anterior": mes_anterior,
                       "clone_annos": clone_annos, "request_from": "poa_cronograma"})

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        try:
            action = request.POST.get('action', '')
            estamento_id = kwargs.get('estamento_id')
            poa_anno = kwargs.get('poa_anno')
            cronograma_id = request.POST.get('id', '')
            cronograma = Cronograma.objects.get(id=cronograma_id)
            cronograma_presupuesto = request.POST.get('cronograma_presupuesto', '')
            evidencia_id = ''
            evidencia_description = ''
            evidencia_url = ''
            evidencia_icon = ''
            evidencia_type = ''
            evidencia_preview = False
            file_name = ''
            index = f'{cronograma.cronograma_actividad.actividad_meta.meta_operativo.operativo_order}.' \
                    f'{cronograma.cronograma_actividad.actividad_meta.meta_order}.' \
                    f'{cronograma.cronograma_actividad.actividad_order}.'

            if action == 'notas':
                cronograma.cronograma_notas = request.POST.get('cronograma_notas', '')
                cronograma.save()
            elif action == 'presupuesto':
                if cronograma_presupuesto != '':
                    cronograma.cronograma_presupuesto = cronograma_presupuesto
                else:
                    cronograma.cronograma_presupuesto = 0
                cronograma.save()
            elif action == 'eliminar':
                evidencia = Evidencia.objects.filter(evidencia_cronograma=cronograma).first()
                if evidencia:
                    RegistrarLog(estamento_id, poa_anno, request.user, 'Delete', f'Se eliminó la evidencia del mes de {cronograma.cronograma_mes.mes_name} de la actividad {index} con id {evidencia.evidencia_cronograma.cronograma_actividad.id}')
                    evidencia.delete()
            elif action == 'evidencia':
                evidencia = Evidencia.objects.filter(evidencia_cronograma=cronograma).first()
                files = request.FILES
                if files:
                    upload = files['evidencia_file']

                    if evidencia is None:
                        evidencia = Evidencia()
                        evidencia.evidencia_cronograma = cronograma
                        evidencia.save()
                        RegistrarLog(estamento_id, poa_anno, request.user, 'Create', f'Se creó una nueva evidencia con id {evidencia.id} del mes de {cronograma.cronograma_mes.mes_name} en la actividad {index}')
                    elif evidencia.evidencia_file:
                        RegistrarLog(estamento_id, poa_anno, request.user, 'Update', f'Se cambió la evidencia con id {evidencia.id} del mes de {cronograma.cronograma_mes.mes_name} en la actividad {index}')
                        if os.path.exists(evidencia.evidencia_file.path):
                            os.remove(evidencia.evidencia_file.path)

                    now = datetime.datetime.now()
                    timestamp = now.strftime("%d%m%Y%H%M%S")
                    fss = FileSystemStorage()
                    file_name = RenameFile(cronograma.cronograma_actividad.actividad_medio.medio_description[:40])
                    file_name += '_' + str(evidencia.id) + '_' + timestamp + os.path.splitext(upload.name)[1]

                    file = fss.save("evidencias/" + file_name, upload)

                    evidencia.evidencia_file = file

                if evidencia:
                    evidencia.evidencia_description = request.POST.get('evidencia_description', '')
                    evidencia.save()

                url_base = GetURL(request)
                evidencia_id = evidencia.id
                evidencia_description = evidencia.evidencia_description
                evidencia_data = GetEvidenciaType(evidencia.id)
                evidencia_type = evidencia_data['evidencia_type']
                evidencia_preview = evidencia_data['evidencia_preview']
                evidencia_icon = f"{url_base}{evidencia_data['evidencia_icon']}"
                evidencia_url = f"{url_base}{evidencia.evidencia_file.url}"

            data = {
                'success': True,
                'cronograma_id': str(cronograma_id),
                'event_action': action,
                'evidencia_id': evidencia_id,
                'evidencia_description': evidencia_description,
                'evidencia_name': file_name,
                'evidencia_url': evidencia_url,
                'evidencia_type': evidencia_type,
                'evidencia_icon': evidencia_icon,
                'evidencia_preview': evidencia_preview,
                'cronograma_presupuesto': cronograma.cronograma_presupuesto,
            }
            return JsonResponse(data=data)
        except Exception as e:
            messages.error(request, F'Ocurrió un error. ' + str(e))

        return redirect("poa_cronograma", estamento_id=estamento_id, poa_anno=poa_anno)


class POAEstado(View):
    estamento_id = 0
    poa_anno = 0

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.estamento_id = kwargs.get('estamento_id')
        self.poa_anno = kwargs.get('poa_anno')
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        estamento_id = self.estamento_id
        poa_anno = self.poa_anno

        poa_estado = int(request.POST["poa_estado"])
        poa_notas = request.POST["poa_notas"]
        poa = POA.objects.filter(id=request.POST["poa_id"]).first()
        poa.poa_estado_id = poa_estado
        poa.poa_notas = poa_notas

        try:
            poa.save()
        except Exception as e:
            messages.error(request, F'Ocurrió un error. ' + str(e))

        if poa_estado > 1:
            template = "basic_email.html"
            estamento = Estamento.objects.get(id=estamento_id)
            notes = poa_notas
            if notes != "":
                notes = "Notas: " + notes

            if poa_estado == 2:
                notificacion_user = estamento.estamento_user
                user_name = estamento.estamento_user.username
                user_mail = estamento.estamento_user.email
                subject = "Su POA ha cambiado de estado"
                message = "Su POA(" + estamento.estamento_name + ") ha pasado a edición. Ahora puede realizar cambios en el mismo y en caso de tener notas, se le sugiere tenerlas en cuenta para las modificaciones"
                estado = 'edición'
            elif poa_estado == 3:
                estamento_sub = Estamento.objects.get(id=estamento.estamento_sub_id)
                notificacion_user = estamento_sub.estamento_user
                user_name = estamento_sub.estamento_user.username
                user_mail = estamento_sub.estamento_user.email
                subject = "Solicitud de revisión"
                message = "Se le solicita revisar el POA(" + estamento.estamento_name + "). Puede agregar notas en caso de ser necesario, para que el responsable pueda realizar los cambios sugeridos"
                estado = 'revisión'
                OrderAll(self.estamento_id, self.poa_anno, request.user, "la solicitud de revisión")
            else:
                notificacion_user = estamento.estamento_user
                user_name = estamento.estamento_user.username
                user_mail = estamento.estamento_user.email
                subject = "Su POA ha sido aprobado"
                message = "Felicidades, su POA(" + estamento.estamento_name + ") se ha aprobado, desde este momento no podrá realizar más cambios en el mismo"
                estado = 'aprobado'

            RegistrarLog(estamento_id, poa_anno, request.user, 'Update', f'Cambió el estado del poa a {estado}')
            firma = GetURL(request) + "/static/images/logo/popin_firma.png"
            try:
                SendMail(user_name, user_mail, subject, template, message, notes, firma)
                SendNotification(request.user, notificacion_user, "ntf", message)
            except Exception:
                messages.error(request, F'No se pudo enviar el correo de notificación')

        return redirect("poa_preview", estamento_id=estamento_id, poa_anno=poa_anno, pag=1)


class POAIncluirSubs(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        action = kwargs.get('action')

        poa = POA.objects.filter(poa_anno=poa_anno, poa_estamento_id=estamento_id).first()
        if poa.poa_include_subs:
            poa.poa_include_subs = False
        else:
            poa.poa_include_subs = True

        try:
            poa.save()
        except Exception as e:
            messages.error(request, F'Ocurrió un error. ' + str(e))

        if action == 'poa_cronograma':
            return redirect(action, estamento_id=estamento_id, poa_anno=poa_anno)
        return redirect(action, estamento_id=estamento_id, poa_anno=poa_anno, pag=1)


class POAClone(View):
    estamento_id = 0
    poa_anno = 0

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.estamento_id = kwargs.get('estamento_id')
        self.poa_anno = kwargs.get('poa_anno')
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        estamento_id = self.estamento_id
        poa_anno = self.poa_anno

        try:
            clone_anno = request.POST["clone_anno"]
            clone_poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
            total_objetivos = 0
            if clone_poa:
                total_objetivos = ObjetivoOperativo.objects.filter(operativo_poa=clone_poa).count()
            if total_objetivos == 0:
                if clone_poa is None:
                    clone_poa = NewPOA(estamento_id, poa_anno, 2, request.user)
                else:
                    clone_poa.poa_estado_id = 2
                    clone_poa.poa_user_modificacion = request.user
                    clone_poa.save()

                poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=clone_anno).first()
                objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id)
                for objetivo in objetivos:
                    clone_objetivo = ObjetivoOperativo()
                    clone_objetivo.operativo_poa = clone_poa
                    clone_objetivo.operativo_order = objetivo.operativo_order
                    clone_objetivo.operativo_selected = objetivo.operativo_selected
                    clone_objetivo.operativo_linea = objetivo.operativo_linea
                    clone_objetivo.operativo_description = objetivo.operativo_description
                    clone_objetivo.save()

                    metas = Meta.objects.filter(meta_operativo_id=objetivo.id)
                    for meta in metas:
                        clone_meta = Meta()
                        clone_meta.meta_operativo = clone_objetivo
                        clone_meta.meta_order = meta.meta_order
                        clone_meta.meta_selected = meta.meta_selected
                        clone_meta.meta_description = meta.meta_description
                        clone_meta.save()

                        actividades = Actividad.objects.filter(actividad_meta_id=meta.id)
                        for actividad in actividades:
                            clone_actividad = Actividad()
                            clone_actividad.actividad_meta = clone_meta
                            clone_actividad.actividad_order = actividad.actividad_order
                            clone_actividad.actividad_description = actividad.actividad_description
                            clone_actividad.actividad_peso = actividad.actividad_peso
                            clone_actividad.actividad_presupuesto = actividad.actividad_presupuesto
                            clone_actividad.actividad_medio_id = actividad.actividad_medio_id
                            clone_actividad.actividad_responsable_id = actividad.actividad_responsable_id
                            clone_actividad.save()

                            cronogramas = Cronograma.objects.filter(cronograma_actividad_id=actividad.id)
                            for cronograma in cronogramas:
                                clone_cronograma = Cronograma()
                                clone_cronograma.cronograma_actividad = clone_actividad
                                clone_cronograma.cronograma_mes = cronograma.cronograma_mes
                                clone_cronograma.save()

                RegistrarLog(estamento_id, poa_anno, request.user, 'Update', f'Se clonó el poa del {poa_anno}')
                return redirect("poa_edit_start", estamento_id=estamento_id, poa_anno=clone_poa.poa_anno)
            else:
                return redirect("poa_preview", estamento_id=estamento_id, poa_anno=poa_anno, pag=1)
        except Exception as e:
            messages.error(request, F'Ocurrió un error. ' + str(e))


class POAEditStart(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        poa = POA.objects.filter(poa_anno=poa_anno, poa_estamento_id=estamento_id).first()
        if poa is None:
            poa = NewPOA(estamento_id, poa_anno, 1, request.user)
        else:
            if poa_anno < datetime.date.today().year or poa.poa_estado_id > 2:
                return redirect("poa_preview", estamento_id=estamento_id, poa_anno=poa_anno, pag=1)

        objetivo = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id).order_by('operativo_order').first()
        if objetivo:
            ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id).update(operativo_selected=False)
            objetivo.operativo_selected = True
            objetivo.save()

            objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id)
            Meta.objects.filter(meta_operativo_id__in=objetivos, meta_order__gt=1).update(meta_selected=False)
            Meta.objects.filter(meta_operativo_id__in=objetivos, meta_order=1).update(meta_selected=True)

        return redirect("poa_edit", estamento_id=estamento_id, poa_anno=poa_anno)


class POAEditUpdate(View):
    selected_id = 0
    selected_name = ""

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.selected_id = kwargs.get('selected_id')
        self.selected_name = kwargs.get('selected_name')
        return super().dispatch(request, *args, **kwargs)

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        selected_id = self.selected_id
        selected_name = self.selected_name
        exists_metas = True
        try:
            if selected_name == "meta":
                meta = Meta.objects.filter(id=selected_id).first()
                Meta.objects.filter(meta_operativo_id=meta.meta_operativo_id).update(meta_selected=False)
                meta.meta_selected = True
                meta.save()

                index = f'{meta.meta_operativo.operativo_order}.' \
                        f'{meta.meta_order}.'

                RegistrarLog(meta.meta_operativo.operativo_poa.poa_estamento_id,
                             meta.meta_operativo.operativo_poa.poa_anno, request.user, 'Selected',
                             f'Se seleccionó la meta {index} con id {meta.id}')
            else:
                objetivo = ObjetivoOperativo.objects.filter(id=selected_id).first()
                ObjetivoOperativo.objects.filter(operativo_poa_id=objetivo.operativo_poa_id).update(
                    operativo_selected=False)
                objetivo.operativo_selected = True
                objetivo.save()
                metas = Meta.objects.filter(meta_operativo_id=selected_id)
                if metas.count() == 0:
                    exists_metas = False

                RegistrarLog(objetivo.operativo_poa.poa_estamento_id, objetivo.operativo_poa.poa_anno, request.user,
                             'Selected', f'Se seleccionó el objetivo {objetivo.operativo_order}. con id {objetivo.id}')
            return JsonResponse({'success': True, 'exists_metas': exists_metas})
        except Exception as e:
            return JsonResponse({'success': False, 'error': 'Valor no encontrado'})


class POAEdit(TemplateView):
    template_name = 'poa_edit.html'
    estamento_id = 0
    poa_anno = 0
    poa_request = None
    send_values = None

    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def dispatch(self, request, *args, **kwargs):
        self.estamento_id = kwargs.get('estamento_id')
        self.poa_anno = kwargs.get('poa_anno')
        self.poa_request = request
        self.send_values = self.get_send_values()

        if isinstance(self.send_values, tuple):
            redirect_url, params = self.send_values
            return redirect(redirect_url, **params)
        elif self.send_values == "home":
            return redirect("home")

        return super().dispatch(request, *args, **kwargs)

    def get_send_values(self):
        poa_anno = self.poa_anno
        estamento_id = self.estamento_id
        user = self.request.user

        send_values = GetPOAList(self.poa_request, user, estamento_id, poa_anno, 0, "order")

        if send_values[0] == "home":
            return "home"
        elif send_values[0] == "poa_edit_start":
            return "poa_edit_start", {"estamento_id": estamento_id, "poa_anno": poa_anno}

        poa = send_values[8]
        if not poa:
            return "poa_card", {"estamento_id": estamento_id, "poa_anno": poa_anno}

        if poa.poa_anno < datetime.date.today().year or poa.poa_estado_id > 2:
            return "poa_preview", {"estamento_id": estamento_id, "poa_anno": poa_anno, "pag": 1}

        return send_values

    # noinspection PyMethodMayBeStatic
    def post(self, request, *args, **kwargs):
        action = str(request.POST.get('action'))
        if action == "order":
            try:
                column = request.POST.get('target_column')
                dragged = request.POST.get('dragged_item')
                position = int(request.POST.get('new_position'))
                data = OrderEdit(column, dragged, position)
                return JsonResponse(data=data)
            except Exception as e:
                messages.error(request, str(e))
        elif action == "order_all":
            try:
                request_from = request.POST.get('request_from')
                OrderAll(self.estamento_id, self.poa_anno, request.user, request_from)
            except Exception as e:
                messages.error(request, str(e))
        elif action == "edit":
            edit_name = str(request.POST.get('name'))
            SaveEdit(request, self.estamento_id, self.poa_anno, edit_name)
        elif action == "delete":
            del_name = str(request.POST.get('name'))
            del_id = str(request.POST.get('id'))
            DeleteEdit(request, del_name, del_id)
        elif action.startswith("load_level"):
            data = {}
            try:
                action = request.POST['action']
                if action == 'load_level2':
                    data = []
                    for objetivo in Objetivo.objects.filter(objetivo_eje_id=request.POST['id']):
                        data.append({'id': objetivo.id, 'name': objetivo.objetivo_description})
                elif action == 'load_level3':
                    data = []
                    for linea in Linea.objects.filter(linea_objetivo_id=request.POST['id']):
                        data.append({'id': linea.id, 'name': linea.linea_description})
                else:
                    data['error'] = 'Ha ocurrido un error'
            except Exception as e:
                data['error'] = str(e)
            return JsonResponse(data, safe=False)

        return redirect("poa_edit", estamento_id=self.estamento_id, poa_anno=self.poa_anno)

    def get_context_data(self, **kwargs):
        poa_anno = self.poa_anno
        estamento_id = self.estamento_id
        send_values = self.send_values

        poa = send_values[8]

        estamento = Estamento.objects.filter(id=estamento_id).first()
        current_objetivo = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id, operativo_selected=True).order_by(
            'operativo_order').first()

        current_meta = None
        if current_objetivo:
            current_meta = Meta.objects.filter(meta_operativo_id=current_objetivo.id, meta_selected=True).order_by(
                'meta_order').first()

        modal_forms = GetModalForms(poa.id, estamento.estamento_user_id)
        clone_annos = getCloneYears(estamento_id)

        context = super().get_context_data(**kwargs)
        context['estamento_id'] = estamento_id
        context['poa_anno'] = poa_anno
        context['estamento_name'] = estamento.estamento_name
        context['estamento_tipo_id'] = estamento.estamento_tipo_id
        context['estamentos_select'] = send_values[1]
        context['estamentos_list'] = send_values[2]
        context['estamentos_options_todos'] = send_values[3]
        context['can_edit'] = send_values[4]
        context['edit_notas'] = send_values[5]
        context['new_years'] = send_values[6]
        context['poa_years'] = send_values[7]
        context['poa'] = poa
        context['meses'] = Mes.objects.all()
        context['current_objetivo'] = current_objetivo
        context['current_meta'] = current_meta
        context['objetivos_list'] = send_values[9]
        context['notas_poa'] = send_values[12]
        context['tablero_list'] = send_values[13]
        context['poa_list'] = send_values[14]
        context['users_roots'] = send_values[15]
        context['paginas'] = send_values[16]
        context['pag'] = 0
        context['modal_forms'] = modal_forms
        context['clone_annos'] = clone_annos
        context['request_from'] = "poa_edit"

        return context


class POAtoDOC(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        doc_type = kwargs.get('doc_type')

        estamento = Estamento.objects.get(id=estamento_id)
        if estamento:
            poas = POA.objects.filter(poa_estamento=estamento)
            if poas.count() == 0:
                return redirect("home")
            if estamento.estamento_user == user:
                is_current_user = True
            else:
                is_current_user = False
        else:
            return redirect("home")

        poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        values = GetPOAList(request, user, estamento_id, poa_anno, 0, "")
        objetivos_list = values[9]

        template = "poa_to_doc.html"
        doc_out = ConvertHTMLtoDOC(request, template, poa_anno, poa, estamento, objetivos_list, is_current_user,
                                   doc_type)

        return FileResponse(open(doc_out['doc_file'], 'rb'), content_type=doc_out['doc_content'])


class EjecutoriaToDOC(View):
    # noinspection PyMethodMayBeStatic
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        doc_type = kwargs.get('doc_type')

        estamento = Estamento.objects.get(id=estamento_id)
        if estamento:
            poas = POA.objects.filter(poa_estamento=estamento)
            if poas.count() == 0:
                return redirect("home")
        else:
            return redirect("home")

        poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        values = GetPOAList(request, user, estamento_id, poa_anno, 0, "")
        objetivos_list = values[9]

        if poa.poa_estamento.estamento_user == user:
            is_current_user = True
        else:
            is_current_user = False

        template = "ejecutoria_to_doc.html"
        doc_out = ConvertHTMLtoDOC(request, template, poa_anno, poa, estamento, objetivos_list, is_current_user,
                                   doc_type)

        return FileResponse(open(doc_out['doc_file'], 'rb'), content_type=doc_out['doc_content'])


class POAProgreso(View):
    @method_decorator(csrf_exempt)
    @method_decorator(login_required)
    def get(self, request, *args, **kwargs):
        user = request.user
        estamento_id = kwargs.get('estamento_id')
        poa_anno = kwargs.get('poa_anno')
        pag = kwargs.get('pag')

        poa_root = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        poas_subs_list = [poa_root]
        send_values = GetPOAList(request, user, estamento_id, poa_anno, pag, "")
        if IncludeSubs(estamento_id, poa_anno):
            poas_subs_list += send_values[10]

        objetivo_operativos_list = []
        found_poa = False
        for item in poas_subs_list:
            poa = None
            if isinstance(item, POA):
                if pag == 1:
                    poa = item
            elif isinstance(item, dict) and isinstance(item['poa'], POA):
                poa = item['poa']

            if poa:
                found_poa = True
                objetivo_operativos = ObjetivoOperativo.objects.filter(operativo_poa=poa)
                objetivos_list = []
                suma_porciento_meta = 0
                suma_porciento_parcial_meta = 0
                total_meta = 0
                total_parcial_meta = 0
                for objetivo_operativo in objetivo_operativos:
                    metas = Meta.objects.filter(meta_operativo=objetivo_operativo)
                    #if metas.count() > 0:
                    objetivo_index = str(objetivo_operativo.operativo_order) + '.'
                    metas_list = []
                    for meta in metas:
                        actividades = Actividad.objects.filter(actividad_meta=meta)
                        #if actividades.count() > 0:
                        meta_index = objetivo_index + str(meta.meta_order) + '.'
                        actividades_list = []
                        total_meta += 1
                        presupuesto = 0
                        porciento_meta = 0
                        porciento_parcial_meta = 0
                        peso_parcial_meta = 0
                        for actividad in actividades:
                            porcientos = CalcularPorcientos(actividad, presupuesto, porciento_meta,
                                                            porciento_parcial_meta, peso_parcial_meta)
                            presupuesto_gastado = porcientos[1]
                            porciento_meta = porcientos[2]
                            porciento_actividad = porcientos[3]
                            porciento_parcial_meta = porcientos[4]
                            porciento_parcial_actividad = porcientos[5]
                            peso_parcial_meta = porcientos[6]
                            cronogramas = Cronograma.objects.filter(cronograma_actividad=actividad)

                            actividad_index = meta_index + str(actividad.actividad_order) + '.'
                            actividades_list.append(
                                {'actividad_index': actividad_index, 'actividad': actividad,
                                 'porciento_actividad': porciento_actividad,
                                 'porciento_parcial_actividad': porciento_parcial_actividad,
                                 'presupuesto_gastado': presupuesto_gastado,
                                 'cronogramas': cronogramas})

                        porciento_meta = GetFormatPorciento(porciento_meta)
                        porciento_parcial_meta = GetFormatPorciento(porciento_parcial_meta)
                        suma_porciento_meta += porciento_meta
                        if peso_parcial_meta:
                            suma_porciento_parcial_meta += porciento_parcial_meta * 100 / peso_parcial_meta
                            suma_porciento_parcial_meta = round(suma_porciento_parcial_meta, 2)
                            total_parcial_meta += 1

                        metas_list.append(
                            {'meta_index': meta_index, 'meta': meta, 'porciento_meta': porciento_meta,
                             'porciento_parcial_meta': porciento_parcial_meta,
                             'actividades_list': actividades_list})

                    objetivos_list.append(
                        {'objetivo_index': objetivo_index, 'objetivo_operativo': objetivo_operativo,
                         'metas_list': metas_list})

                if total_meta > 0:
                    poa_porciento = suma_porciento_meta / total_meta
                    poa_porciento = round(poa_porciento, 2)
                else:
                    poa_porciento = 0
                if total_parcial_meta > 0:
                    poa_porciento_parcial = suma_porciento_parcial_meta / total_parcial_meta
                    poa_porciento_parcial = round(poa_porciento_parcial, 2)
                else:
                    poa_porciento_parcial = 0

                poa_porciento = GetFormatPorciento(poa_porciento)
                poa_porciento_parcial = GetFormatPorciento(poa_porciento_parcial)
                objetivo_operativos_list.append(
                    {'estamento_name': poa.poa_estamento.estamento_name, 'poa_porciento': f'{poa_porciento}%',
                     'poa_porciento_parcial': f'{poa_porciento_parcial}%', 'objetivos_list': objetivos_list})

        if not found_poa:
            return redirect('poa_card', estamento_id=estamento_id, poa_anno=poa_anno, pag=1)

        clone_annos = getCloneYears(estamento_id)

        url_base = GetURL(request)
        return render(request, "poa_progreso.html",
                      {"estamento_id": estamento_id, "poa_anno": poa_anno, "poa": poa_root,
                       "objetivo_operativos_list": objetivo_operativos_list,
                       "url_base": url_base, "estamentos_select": send_values[1],
                       "estamentos_list": send_values[2], "estamentos_options_todos": send_values[3],
                       "can_edit": send_values[4], "edit_notas": send_values[5], "new_years": send_values[6],
                       "poa_years": send_values[7], "poa": send_values[8],
                       "objetivos_list": send_values[9], "notas_poa": send_values[12], "tablero_list": send_values[13],
                       "poa_list": send_values[14], "users_roots": send_values[15], "paginas": send_values[16],
                       "pag": pag, "clone_annos": clone_annos, "request_from": "poa_progreso"})


def getCloneYears(estamento_id):
    if datetime.date.today().month >= MonthNewPoa():
        poas = POA.objects.filter(poa_estamento_id=estamento_id).exclude(
            poa_anno__gt=datetime.date.today().year).order_by(
            "-poa_anno")
    else:
        poas = POA.objects.filter(poa_estamento_id=estamento_id).exclude(
            poa_anno__gte=datetime.date.today().year).order_by(
            "-poa_anno")
    poa_annos = poas.values_list('poa_anno', flat=True).distinct()
    return poa_annos


def NewPOA(estamento_id, poa_anno, poa_estado_id, user):
    poa = POA()
    poa.poa_estamento_id = estamento_id
    poa.poa_anno = poa_anno
    poa.poa_estado_id = poa_estado_id
    poa.poa_user_modificacion = user
    poa.save()
    return poa


def OrderEdit(column, dragged, position):
    updated_values = []
    if column.startswith('meta_actividad_'):
        if position > 0:
            column_id = int(column.rsplit('_', 1)[1])
            dragged_id = int(dragged.rsplit('_', 1)[1])
            meta_id = column_id
            actividad_id = dragged_id
            actividad = Actividad.objects.filter(id=actividad_id).first()

            if actividad.actividad_order > position:
                index = position + 1
                actividades = Actividad.objects.filter(
                    actividad_meta_id=meta_id,
                    actividad_order__gte=position
                ).exclude(id=actividad_id)
            else:
                index = 1
                actividades = Actividad.objects.filter(
                    actividad_meta_id=meta_id,
                    actividad_order__lte=position
                ).exclude(id=actividad_id)

            for act in actividades:
                act.actividad_order = index
                act.save()
                index += 1

            actividad.actividad_order = position
            actividad.save()

            updated_actividades = Actividad.objects.filter(actividad_meta_id=meta_id).order_by(
                'actividad_order')
            updated_values = [
                {"id": actividad.id,
                 "value": f'{actividad.actividad_order}.'} for actividad in updated_actividades]
    else:
        if column.startswith('add_meta_actividad_'):
            column_id = int(column.rsplit('_', 1)[1])
            dragged_id = int(dragged.rsplit('_', 1)[1])
            actividad_id = dragged_id
            actividad = Actividad.objects.filter(id=actividad_id).first()
            meta_id = actividad.actividad_meta_id
            if column_id != meta_id:
                actividad_order = actividad.actividad_order
                Actividad.objects.filter(
                    actividad_meta_id=meta_id,
                    actividad_order__gt=actividad_order
                ).update(actividad_order=F('actividad_order') - 1)
                actividad.actividad_meta_id = column_id
                actividad.actividad_order = position
                actividad.save()
        else:
            if column.startswith('objetivo_meta_'):
                updated_values = []
                if position > 0:
                    column_id = int(column.rsplit('_', 1)[1])
                    dragged_id = int(dragged.rsplit('_', 1)[1])
                    objetivo_id = column_id
                    meta_id = dragged_id
                    meta = Meta.objects.filter(id=meta_id).first()

                    if meta.meta_order > position:
                        index = position + 1
                        metas = Meta.objects.filter(
                            meta_operativo_id=objetivo_id,
                            meta_order__gte=position
                        ).exclude(id=meta_id)
                    else:
                        index = 1
                        metas = Meta.objects.filter(
                            meta_operativo_id=objetivo_id,
                            meta_order__gt=meta.meta_order,
                            meta_order__lte=position
                        ).exclude(id=meta_id)

                    for met in metas:
                        met.meta_order = index
                        met.save()
                        index += 1

                    meta.meta_order = position
                    meta.save()

                    updated_metas = Meta.objects.filter(meta_operativo_id=objetivo_id).order_by(
                        'meta_order')
                    updated_values = [{"id": meta.id,
                                       "value": f'{meta.meta_order}.'} for meta in updated_metas]
            else:
                if column.startswith('add_objetivo_meta_'):
                    column_id = int(column.rsplit('_', 1)[1])
                    dragged_id = int(dragged.rsplit('_', 1)[1])
                    meta_id = dragged_id
                    meta = Meta.objects.filter(id=meta_id).first()
                    objetivo_id = meta.meta_operativo_id
                    if column_id != objetivo_id:
                        meta_order = meta.meta_order
                        Meta.objects.filter(
                            meta_operativo_id=objetivo_id,
                            meta_order__gt=meta_order
                        ).update(meta_order=F('meta_order') - 1)
                        if meta.meta_selected:
                            Meta.objects.filter(
                                meta_operativo_id=objetivo_id, meta_order=1
                            ).update(meta_selected=True)

                        meta.meta_operativo_id = column_id
                        meta.meta_order = position
                        meta.meta_selected = False
                        meta.save()
                else:
                    if column.startswith('poa_objetivo_'):
                        dragged_id = int(dragged.rsplit('_', 1)[1])
                        objetivo_id = dragged_id
                        objetivo = ObjetivoOperativo.objects.filter(id=objetivo_id).first()

                        if objetivo.operativo_order > position:
                            index = position + 1
                            objetivos = ObjetivoOperativo.objects.filter(
                                operativo_poa_id=objetivo.operativo_poa_id,
                                operativo_order__gte=position
                            ).exclude(id=objetivo_id)
                        else:
                            index = 1
                            objetivos = ObjetivoOperativo.objects.filter(
                                operativo_poa_id=objetivo.operativo_poa_id,
                                operativo_order__lte=position
                            ).exclude(id=objetivo_id)

                        for obj in objetivos:
                            obj.operativo_order = index
                            obj.save()
                            index += 1

                        objetivo.operativo_order = position
                        objetivo.save()

                        updated_objetivos = ObjetivoOperativo.objects.filter(
                            operativo_poa_id=objetivo.operativo_poa_id).order_by('operativo_order')
                        updated_values = [{"id": objetivo.id,
                                           "value": f'{objetivo.operativo_order}.'} for objetivo in
                                          updated_objetivos]

    return {"updated_values": updated_values}


def SaveEdit(request, estamento_id, poa_anno, action):
    try:
        user = request.user
        log_action = 'Update'
        log_description = ''
        if action == "poa":
            form = FormPOA(request.POST)
            if form.is_valid():
                poa = POA.objects.filter(id=request.POST["poa_id"]).first()
                if poa is None:
                    poa = POA()
                    poa.poa_anno = poa_anno
                    poa.poa_estamento_id = request.POST["poa_estamento"]
                    poa.poa_estado_id = 1
                    log_action = 'Create'
                    log_description = f'Se creó un nuevo poa con id '
                else:
                    poa.poa_estado_id = 2
                    log_description = f'Se actualizó el poa con id '
                poa.poa_user_modificacion = user
                poa.save()
                log_description += f'{poa.id}'
            else:
                messages.error(request, F"Error guardando")
        elif action == "objetivo":
            form = FormObjetivoOperativo(request.POST)
            if form.is_valid():
                poa = POA.objects.filter(id=request.POST["poa_id"]).first()
                operativo_poa = poa
                operativo_linea = Linea.objects.filter(id=request.POST["level3"]).first()
                operativo_description = request.POST["operativo_description"]
                objetivo = ObjetivoOperativo.objects.filter(id=request.POST["operativo_id"]).first()
                objetivo_exists = None
                if objetivo is None:
                    objetivo = ObjetivoOperativo()
                    objetivos = ObjetivoOperativo.objects.filter(operativo_poa=operativo_poa)
                    order = len(objetivos) + 1
                    objetivo.operativo_order = order
                    if order == 1:
                        objetivo.operativo_selected = True
                        if poa.poa_estado_id == 1:
                            poa.poa_estado_id = 2
                            poa.save()
                    log_action = 'Create'
                    log_description = f'Se creó un nuevo objetivo {objetivo.operativo_order}. con id '

                    objetivo_exists = ObjetivoOperativo.objects.filter(operativo_poa=operativo_poa,
                                                                       operativo_linea=operativo_linea,
                                                                       operativo_description=operativo_description).first()
                else:
                    log_description = f'Se actualizó el objetivo {objetivo.operativo_order}. con id '

                if objetivo_exists is None:
                    objetivo.operativo_poa = operativo_poa
                    objetivo.operativo_linea = operativo_linea
                    objetivo.operativo_description = operativo_description
                    objetivo.save()
                    log_description += f'{objetivo.id}'
                else:
                    log_description = f'Se intentó crear un objetivo con la misma descripción {operativo_description}'
                    messages.error(request, F"Ya existe el operativo operativo")
            else:
                for field, items in form.errors.items():
                    for item in items:
                        messages.error(request, '{}: {}'.format(field, item))
        elif action == "meta":
            form = FormMeta(request.POST)
            if form.is_valid():
                poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
                objetivo = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id, operativo_selected=True).order_by(
                    'operativo_order').first()
                meta = Meta.objects.filter(id=request.POST["meta_id"]).first()
                if meta is None:
                    meta = Meta()
                    metas = Meta.objects.filter(meta_operativo=objetivo)
                    order = len(metas) + 1
                    meta.meta_order = order
                    if order == 1:
                        meta.meta_selected = True
                    log_action = 'Create'
                    log_description = f'Se creó una nueva meta {objetivo.operativo_order}.{meta.meta_order}. con id '
                else:
                    log_description = f'Se actualizó la meta {objetivo.operativo_order}.{meta.meta_order}. con id '

                meta.meta_operativo = objetivo
                meta.meta_description = request.POST["meta_description"]
                meta.save()
                log_description += f'{meta.id}'
            else:
                for field, items in form.errors.items():
                    for item in items:
                        messages.error(request, '{}: {}'.format(field, item))
        elif action == "actividad":
            form = FormActividad(request.POST)
            if form.is_valid():
                poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
                objetivo = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id, operativo_selected=True).order_by(
                    'operativo_order').first()
                meta = Meta.objects.filter(meta_operativo_id=objetivo.id, meta_selected=True).order_by(
                    'meta_order').first()
                actividad = Actividad.objects.filter(id=request.POST.get("actividad_id")).first()
                if actividad is None:
                    actividad = Actividad()
                    actividades = Actividad.objects.filter(actividad_meta=meta)
                    actividad.actividad_order = len(actividades) + 1
                    log_action = 'Create'
                    log_description = f'Se creó una nueva actividad {meta.meta_order}.{actividad.actividad_order}. con id '
                else:
                    log_description = f'Se actualizó la actividad {meta.meta_order}.{actividad.actividad_order}. con id '

                actividad.actividad_meta = meta
                actividad.actividad_description = request.POST.get("actividad_description")
                actividad_peso = request.POST.get("actividad_peso", '')
                if actividad_peso == '':
                    actividad_peso = 0
                actividad.actividad_peso = actividad_peso
                actividad_presupuesto = request.POST.get("actividad_presupuesto", '')
                if actividad_presupuesto == '':
                    actividad_presupuesto = 0
                actividad.actividad_presupuesto = actividad_presupuesto

                actividad_medio = request.POST.get("actividad_medio_nuevo")
                if actividad_medio == '':
                    actividad_medio = request.POST.get("actividad_medio")
                    medio = MedioVerificacion.objects.filter(id=actividad_medio).first()
                else:
                    medio = MedioVerificacion.objects.filter(medio_description=actividad_medio,
                                                             medio_user_id=user.id).first()
                    if medio is None:
                        medio = MedioVerificacion()
                        medio.medio_description = actividad_medio
                        medio.medio_user = user
                        medio.save()
                actividad.actividad_medio = medio

                actividad_responsable = request.POST.get("actividad_responsable_nuevo")
                if actividad_responsable == '':
                    actividad_responsable = request.POST.get("actividad_responsable")
                    responsable = Responsable.objects.filter(id=actividad_responsable).first()
                else:
                    responsable = Responsable.objects.filter(responsable_description=actividad_responsable,
                                                             responsable_user_id=user.id).first()
                    if responsable is None:
                        responsable = Responsable()
                        responsable.responsable_description = actividad_responsable
                        responsable.responsable_user = user
                        responsable.save()
                actividad.actividad_responsable = responsable
                actividad.save()
                log_description += f'{actividad.id}'

                poa = actividad.actividad_meta.meta_operativo.operativo_poa
                if poa.poa_estado.id == 1:
                    poa.poa_estado_id = 2
                    poa.save()

                meses = request.POST.getlist("actividad_cronograma")
                cronogramas = Cronograma.objects.filter(cronograma_actividad=actividad).exclude(
                    cronograma_mes__in=meses)
                for cronograma in cronogramas:
                    cronograma.delete()

                cronogramas = Cronograma.objects.filter(cronograma_actividad=actividad)
                cronograma_meses = [str(cronograma.cronograma_mes_id) for cronograma in cronogramas]
                meses_a_crear = [mes for mes in meses if mes not in cronograma_meses]
                for mes_id in meses_a_crear:
                    mes = Mes.objects.filter(id=mes_id).first()
                    cronograma = Cronograma()
                    cronograma.cronograma_mes = mes
                    cronograma.cronograma_cumplimiento_mes = mes
                    cronograma.cronograma_actividad = actividad
                    cronograma.save()
            else:
                for field, items in form.errors.items():
                    for item in items:
                        messages.error(request, '{}: {}'.format(field, item))
        else:
            return redirect("home")

        RegistrarLog(estamento_id, poa_anno, request.user, log_action, log_description)
    except Exception as e:
        messages.error(request, F'Ocurrió un error. ' + str(e))


def DeleteEdit(request, action, del_id):
    try:
        estamento_id = 0
        poa_anno = 0
        if action == "poa":
            poa = POA.objects.get(id=del_id)
            estamento_id = poa.poa_estamento_id
            poa_anno = poa.poa_anno
            log_description = f'Se eliminó el poa {del_id}'
            DeletePOA(del_id)
        elif action == "objetivo":
            objetivo = ObjetivoOperativo.objects.get(id=del_id)
            estamento_id = objetivo.operativo_poa.poa_estamento_id
            poa_anno = objetivo.operativo_poa.poa_anno
            log_description = f'Se eliminó el objetivo {objetivo.operativo_order}. con id {del_id}'
            DeleteObjetivoOperativo(del_id)
        elif action == "meta":
            meta = Meta.objects.get(id=del_id)
            estamento_id = meta.meta_operativo.operativo_poa.poa_estamento_id
            poa_anno = meta.meta_operativo.operativo_poa.poa_anno
            log_description = f'Se eliminó la meta {meta.meta_operativo.operativo_order}.{meta.meta_order}. con id {del_id}'
            DeleteMeta(del_id)
        elif action == "actividad":
            actividad = Actividad.objects.get(id=del_id)
            estamento_id = actividad.actividad_meta.meta_operativo.operativo_poa.poa_estamento_id
            poa_anno = actividad.actividad_meta.meta_operativo.operativo_poa.poa_anno
            log_description = f'Se eliminó la actividad {actividad.actividad_meta.meta_operativo.operativo_order}.{actividad.actividad_meta.meta_order}.{actividad.actividad_order}. con id {del_id}'
            DeleteActividad(del_id)
        RegistrarLog(estamento_id, poa_anno, request.user, "Delete", log_description)
    except Exception as e:
        messages.error(request, F'Ocurrió un error. ' + str(e))


def GetModalForms(poa_id, user_id):
    objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=poa_id)

    list_forms_operativos = []
    for objetivo in objetivos:
        form_operativo = FormObjetivoOperativo(
            initial={'poa_id': poa_id, 'operativo_id': objetivo.id,
                     'linea_id': objetivo.operativo_linea.id})
        form_operativo.fields['level1'].initial = objetivo.operativo_linea.linea_objetivo.objetivo_eje_id
        form_operativo.fields['level2'].queryset = Objetivo.objects.filter(
            objetivo_eje_id=objetivo.operativo_linea.linea_objetivo.objetivo_eje_id)
        form_operativo.fields['level2'].initial = objetivo.operativo_linea.linea_objetivo_id
        form_operativo.fields['level3'].queryset = Linea.objects.filter(
            linea_objetivo_id=objetivo.operativo_linea.linea_objetivo_id)
        form_operativo.fields['level3'].initial = objetivo.operativo_linea_id
        form_operativo.fields['operativo_description'].initial = objetivo.operativo_description

        metas = Meta.objects.filter(meta_operativo=objetivo)
        list_forms_metas = []
        for meta in metas:
            actividades = Actividad.objects.filter(actividad_meta=meta)
            list_forms_actividades = []
            peso_total = 0
            for actividad in actividades:
                form_actividad = FormActividad(
                    initial={'poa_id': poa_id,
                             'actividad_id': actividad.id, 'meta_id': meta.id})
                form_actividad.fields['actividad_description'].initial = actividad.actividad_description
                form_actividad.fields['actividad_presupuesto'].initial = actividad.actividad_presupuesto
                form_actividad.fields['actividad_peso'].initial = actividad.actividad_peso
                peso_total += actividad.actividad_peso

                form_actividad.fields['actividad_medio'].queryset = MedioVerificacion.objects.filter(
                    Q(medio_user_id=user_id) | Q(id=actividad.actividad_medio_id))
                form_actividad.fields['actividad_medio'].initial = actividad.actividad_medio_id

                form_actividad.fields['actividad_responsable'].queryset = Responsable.objects.filter(
                    Q(responsable_user_id=user_id) | Q(id=actividad.actividad_responsable_id))
                form_actividad.fields['actividad_responsable'].initial = actividad.actividad_responsable_id

                actividad_cronograma = [cronograma for cronograma in Cronograma.objects.filter(
                    cronograma_actividad=actividad).values_list("cronograma_mes", flat=True)]
                form_actividad.fields['actividad_cronograma'].initial = actividad_cronograma

                list_forms_actividades.append(
                    {'actividad': actividad, 'cronograma': actividad_cronograma, 'form_actividad': form_actividad})

            actividad_peso = 0
            if peso_total <= 100:
                actividad_peso = 100 - peso_total

            last_actividad = Actividad.objects.all().last()
            if last_actividad:
                actividad_medio_id = last_actividad.actividad_medio_id
                actividad_responsable_id = last_actividad.actividad_responsable_id

            list_actividades = {'actividad_peso': actividad_peso,
                                'actividad_medio_id': actividad_medio_id,
                                'actividad_responsable_id': actividad_responsable_id,
                                'list_forms_actividades': list_forms_actividades}
            form_meta = FormMeta(
                initial={'poa_id': poa_id, 'meta_id': meta.id, 'operativo_id': objetivo.id})
            form_meta.fields['meta_description'].initial = meta.meta_description

            peso_error = ""
            if peso_total != 100:
                peso_error = "La sumatoria(" + str(
                    peso_total) + "%) de los pesos de todas las actividades de la meta debe ser 100%"
            list_forms_metas.append(
                {'meta': meta, 'form_meta': form_meta, 'list_actividades': list_actividades,
                 'peso_error': peso_error})

        list_forms_operativos.append({'objetivo': objetivo, 'form_operativo': form_operativo,
                                      'list_forms_metas': list_forms_metas})

    form_operativo_add = FormObjetivoOperativo(initial={'poa_id': poa_id, 'operativo_id': 0})
    form_operativo_add.fields['level1'].initial = 0
    form_operativo_add.fields['level2'].queryset = Objetivo.objects.filter(objetivo_eje_id=1)
    form_operativo_add.fields['level2'].initial = 0
    form_operativo_add.fields['level3'].queryset = Linea.objects.filter(linea_objetivo_id=1)
    form_operativo_add.fields['level3'].initial = 0

    form_meta_add = FormMeta()

    form_actividad_add = FormActividad()
    form_actividad_add.fields['actividad_medio'].queryset = MedioVerificacion.objects.filter(
        medio_user_id=user_id)
    form_actividad_add.fields['actividad_responsable'].queryset = Responsable.objects.filter(
        responsable_user_id=user_id)
    last_actividad = Actividad.objects.all().last()
    if last_actividad:
        actividad_medio_id = last_actividad.actividad_medio_id
        actividad_responsable_id = last_actividad.actividad_responsable_id
        form_actividad_add.fields['actividad_medio'].initial = actividad_medio_id
        form_actividad_add.fields['actividad_responsable'].initial = actividad_responsable_id

    return {'form_operativo_add': form_operativo_add, 'form_meta_add': form_meta_add,
            'form_actividad_add': form_actividad_add, 'list_forms_operativos': list_forms_operativos}


def DeletePOA(poa_id):
    try:
        poa = POA.objects.get(id=poa_id)
        if poa:
            objetivo_operativo_list = ObjetivoOperativo.objects.filter(operativo_poa_id=poa.id)
            for objetivo_operativo in objetivo_operativo_list:
                DeleteObjetivoOperativo(objetivo_operativo.id)

            poa.delete()
    except Exception as e:
        raise Exception(str(e))


def DeleteObjetivoOperativo(objetivo_id):
    try:
        objetivo = ObjetivoOperativo.objects.get(id=objetivo_id)
        if objetivo:
            poa_id = objetivo.operativo_poa.id
            total_objetivos = ObjetivoOperativo.objects.filter(operativo_poa_id=objetivo.operativo_poa.id).count()
            meta_list = Meta.objects.filter(meta_operativo_id=objetivo.id)
            for meta in meta_list:
                DeleteMeta(meta.id)

            operativo_list = ObjetivoOperativo.objects.filter(
                operativo_poa_id=objetivo.operativo_poa.id,
                operativo_order__gt=objetivo.operativo_order
            )
            for operativo in operativo_list:
                operativo.operativo_order -= 1
                operativo.save()

            if objetivo.operativo_selected:
                first_objetivo = ObjetivoOperativo.objects.filter(operativo_poa_id=objetivo.operativo_poa_id).order_by(
                    'operativo_order').first()
                first_objetivo.operativo_selected = True
                first_objetivo.save()
            objetivo.delete()

            if total_objetivos == 1:
                poa = POA.objects.get(id=poa_id)
                poa.poa_estado_id = 1
                poa.save()
    except Exception as e:
        raise Exception(str(e))


def DeleteMeta(meta_id):
    try:
        meta = Meta.objects.get(id=meta_id)
        if meta:
            actividad_list = Actividad.objects.filter(actividad_meta_id=meta.id)
            for actividad in actividad_list:
                DeleteActividad(actividad.id)

            meta_list = Meta.objects.filter(
                meta_operativo=meta.meta_operativo.id,
                meta_order__gt=meta.meta_order
            )
            for meta_item in meta_list:
                meta_item.meta_order -= 1
                meta_item.save()

            if meta.meta_selected:
                first_meta = Meta.objects.filter(meta_operativo_id=meta.meta_operativo_id).order_by(
                    'meta_order').first()
                first_meta.meta_selected = True
                first_meta.save()

            meta.delete()
    except Exception as e:
        raise Exception(str(e))


def DeleteActividad(actividad_id):
    try:
        actividad = Actividad.objects.get(id=actividad_id)
        if actividad:
            cronograma_list = Cronograma.objects.filter(cronograma_actividad_id=actividad.id)
            for cronograma in cronograma_list:
                DeleteCronograma(cronograma.id)

            actividad_list = Actividad.objects.filter(
                actividad_meta=actividad.actividad_meta.id,
                actividad_order__gt=actividad.actividad_order
            )
            for meta_id in actividad_list:
                meta_id.actividad_order -= 1
                meta_id.save()

            actividad.delete()
    except Exception as e:
        raise Exception(str(e))


def DeleteCronograma(cronograma_id):
    try:
        cronograma = Cronograma.objects.get(id=cronograma_id)
        if cronograma:
            Evidencia.objects.filter(evidencia_cronograma_id=cronograma_id).delete()
            cronograma.delete()
    except Exception as e:
        raise Exception(str(e))


def EstamentosSortList(estamentos_list):
    sort_estamentos_list = []
    estamento_level_list = []
    estamento_id_list = []
    for estamento in estamentos_list:
        if estamento.id not in estamento_id_list:
            estamento_id_list.append(estamento.id)
            level = EstamentoLevel(estamento.id)
            estamento_level_list.append({"level": level, "estamento": estamento, "tipo": estamento.estamento_tipo_id})

    estamento_level_list.sort(key=GetTipo)
    for estamento_level in estamento_level_list:
        sort_estamentos_list.append(estamento_level["estamento"])

    return sort_estamentos_list


def EstamentoLevel(estamento_id):
    level = 0
    estamento = Estamento.objects.get(id=estamento_id)
    if estamento:
        if estamento.estamento_sub_id:
            level = EstamentoLevel(estamento.estamento_sub_id) + 1

    return level


def GetLevel(level):
    return level['level']


def GetTipo(tipo):
    return tipo['tipo']


def EstamentosAdminList():
    exclude_list = list(Estamento.objects.filter(estamento_has_poa=True))
    estamentos = Estamento.objects.filter(estamento_has_poa=False)
    for estamento in estamentos:
        if estamento.estamento_has_poa:
            exclude_list.append(estamento)
        else:
            if IsRoot(estamento.id):
                exclude_list.append(estamento)

    return exclude_list


def EstamentosSubsList(user, estamento_id, estamentos_list, poa_estamentos_list, poa_anno):
    estamento_sub_list = []
    poas_sub_list = []
    tableros_sub_list = []
    for estamentos_sub in estamentos_list:
        if estamentos_sub.estamento_main or (not estamentos_sub.estamento_main and estamentos_sub.id == estamento_id):
            estamento_sub_add = EstamentosSubs(estamentos_sub.id)
            if len(estamento_sub_add) > 0:
                estamento_sub_list += estamento_sub_add

    for estamentos_sub in poa_estamentos_list:
        if estamentos_sub.estamento_main or (not estamentos_sub.estamento_main and estamentos_sub.id == estamento_id):
            sub_add = POAEstamentosSubs(user, estamento_id, estamentos_sub.id, poa_anno)
            if len(sub_add) > 0:
                poas_sub_list += sub_add[0]
                tableros_sub_list += sub_add[1]

    return [estamento_sub_list, poas_sub_list, tableros_sub_list]


def EstamentosSubs(estamento_sub_id):
    estamento_subs = Estamento.objects.filter(estamento_sub_id=estamento_sub_id)
    estamento_sub_list = []
    for estamento_sub in estamento_subs:
        if estamento_sub.estamento_has_poa or IsRoot(estamento_sub.id):
            estamento_sub_list.append(estamento_sub)
            estamento_sub_add = EstamentosSubs(estamento_sub.id)
            if len(estamento_sub_add) > 0:
                estamento_sub_list += estamento_sub_add

    return estamento_sub_list


def POAEstamentosSubs(user, root_estamento_id, estamento_sub_id, poa_anno):
    estamento_subs = Estamento.objects.filter(estamento_sub_id=estamento_sub_id)
    poas_sub_list = []
    tableros_sub_list = []
    for estamento_sub in estamento_subs:
        if estamento_sub.estamento_has_poa or IsRoot(estamento_sub.id):
            poas = POA.objects.filter(poa_estamento_id=estamento_sub.id)
            if poas.count() > 0:
                poa_color = GetChildColor(estamento_sub, user)
                poa = POA.objects.filter(poa_estamento_id=estamento_sub.id, poa_anno=poa_anno).first()
                if poa:
                    total_objetivos = ObjetivoOperativo.objects.filter(operativo_poa=poa).count()
                    poa = POA.objects.filter(poa_estamento_id=estamento_sub.id, poa_anno=poa_anno).annotate(
                        poa_color=Value(poa_color, output_field=CharField()),
                        poa_objetivos=Value(total_objetivos, output_field=IntegerField()),
                        poa_user_owner=Value(IsPOAUser(estamento_sub.id, user), output_field=BooleanField()),
                    ).first()

                    notas_poa = ''
                    if poa.poa_estado_id > 1:
                        notas_poa = GetNotasPoa(poa.id)
                        tableros_sub_list = list(
                            TableauEstamento.objects.filter(tablero_estamento=estamento_sub,
                                                            tablero_tableau__tableau_anno=poa_anno))
                    users_roots = getUsersRoots(poa.poa_estamento, user)
                    poas_sub_list.append(
                        {"poa": poa, "notas_poa": notas_poa, 'poa_porciento': '0%', 'poa_porciento_parcial': '0%', 'users_roots': users_roots})
                else:
                    poa = POA()
                    poa.id = 0
                    poa.poa_estamento = estamento_sub
                    poa.poa_estado_id = 0
                    poa.poa_color = '#B7BBCA'
                    poa.poa_user_owner = IsPOAUser(estamento_sub.id, user)
                    poa.poa_objetivos = 0
                    users_roots = getUsersRoots(poa.poa_estamento, user)
                    poas_sub_list.append(
                        {'poa': poa, 'notas_poa': '', 'poa_porciento': '0%', 'poa_porciento_parcial': '0%', 'users_roots': users_roots})

                sub_add = POAEstamentosSubs(user, root_estamento_id, estamento_sub.id, poa_anno)
                if len(sub_add) > 0:
                    poas_sub_list += sub_add[0]
                    tableros_sub_list += sub_add[1]

    return [poas_sub_list, tableros_sub_list]


def GetPOAs(user, estamento_id, poa_anno, pag, filtred_by):
    poa_redirect = ""
    estamento = Estamento.objects.get(id=estamento_id)
    estamentos_select = None
    estamentos_list = None
    estamentos_options_todos = ""
    objetivos_list = None
    poa = None
    new_years = None
    poa_years = None
    can_edit = False
    edit_notas = False
    poas_list = []
    poas_subs_list = []
    cronogramas_list = []
    tablero_list = []
    notas_poa = ""
    users_roots = []
    paginas = []
    paginado = 0

    if estamento is None:
        poa_redirect = "home"
    else:
        poas = POA.objects.filter(poa_estamento=estamento)
        poa_color = GetColor(estamento, user)
        poa = POA.objects.filter(poa_estamento=estamento, poa_anno=poa_anno).first()
        total_objetivos = 0
        if poa:
            total_objetivos = ObjetivoOperativo.objects.filter(operativo_poa=poa).count()
        poa = POA.objects.filter(poa_estamento=estamento, poa_anno=poa_anno).annotate(
            poa_color=Value(poa_color, output_field=CharField()),
            poa_objetivos=Value(total_objetivos, output_field=IntegerField()),
            poa_user_owner=Value(IsPOAUser(estamento_id, user), output_field=BooleanField()),
        ).first()

        if estamento.estamento_user == user and estamento.estamento_has_poa and (poas.count() == 0 or poa is None):
            poa_redirect = "poa_edit_start"
        else:
            values = GetPOAsLists(user, estamento, poa_anno, poa, poas, pag)
            estamentos_select = values[0]
            estamentos_list = values[1]
            poas_subs_list = values[2]
            tableros_sub_list = values[3]
            estamentos_options_todos = values[4]
            can_edit = values[5]
            edit_notas = values[6]
            new_years = values[7]
            poa_years = values[8]
            paginas = list(range(1, values[9] + 1))
            paginado = values[10]

            if IsRoot(estamento_id) or poa:
                if filtred_by == "card":
                    objetivos = ObjetivosList().get_objetivos_list(poa, user, estamento_id, poa_anno, poas_subs_list,
                                                                   tableros_sub_list, pag, filtred_by)
                    poas_list = objetivos['poas_list']
                if filtred_by == "tablero":
                    objetivos = ObjetivosList().get_objetivos_list(poa, user, estamento_id, poa_anno, poas_subs_list,
                                                                   tableros_sub_list, pag, filtred_by)
                    tablero_list = objetivos['tableros_list']
                else:
                    if filtred_by == "" or filtred_by == "order" or filtred_by == "actividades":
                        objetivos = ObjetivosList().get_objetivos_list(poa, user, estamento_id, poa_anno,
                                                                       poas_subs_list,
                                                                       tableros_sub_list, pag, filtred_by)
                        objetivos_list = objetivos['objetivos_poa_list']

                        if filtred_by == "actividades":
                            actividades_list = objetivos['actividades_list']

                            for i in range(1, 13):
                                cronogramas_tmp = Cronograma.objects.filter(cronograma_actividad__in=actividades_list,
                                                                            cronograma_mes=i)
                                cronograma_actividades = {}
                                for cronograma in cronogramas_tmp:
                                    index = f'{cronograma.cronograma_actividad.actividad_meta.meta_operativo.operativo_order}' \
                                            f'.{cronograma.cronograma_actividad.actividad_meta.meta_order}' \
                                            f'.{cronograma.cronograma_actividad.actividad_order}'
                                    cronograma_actividades[index] = cronograma.cronograma_actividad_id

                                cronogramas = []
                                planificadas = []
                                realizadas = []
                                sorted_keys = sorted(cronograma_actividades.keys(),
                                                     key=lambda x: tuple(map(int, x.split('.'))))
                                for index in sorted_keys:
                                    actividad_id = cronograma_actividades[index]
                                    cronograma = Cronograma.objects.filter(cronograma_actividad_id=actividad_id,
                                                                           cronograma_mes_id=i).annotate(
                                        cronograma_estado=Case(
                                            When(cronograma_cumplimiento=True, then=Value('realizada')),
                                            When(cronograma_mes_id__lt=datetime.date.today().month,
                                                 then=Value('atrasada')),
                                            default=Value('planificada'),
                                            output_field=CharField()
                                        ),
                                        cronograma_color=Case(
                                            When(cronograma_cumplimiento=True, cronograma_cumplimiento_mes_id=i,
                                                 then=Value('realizada')),
                                            When(cronograma_cumplimiento=True, then=Value('desfasada')),
                                            When(cronograma_mes_id__lt=datetime.date.today().month,
                                                 then=Value('atrasada')),
                                            default=Value('planificada'),
                                            output_field=CharField()
                                        )
                                    ).first()

                                    cronogramas.append(cronograma)
                                    if cronograma.cronograma_cumplimiento:
                                        realizadas.append(cronograma)
                                    else:
                                        planificadas.append(cronograma)

                                if poa and not poa.poa_include_subs:
                                    planificadas = cronogramas
                                    realizadas = cronogramas

                                mes = Mes.objects.get(id=i)
                                cronogramas_list.append(
                                    {"cronogramas": cronogramas, "planificadas": planificadas, "realizadas": realizadas,
                                     "mes": mes, "anno": poa_anno})
                    else:
                        objetivos_list = ObjetivosList().get_ejes_list(poa, user, estamento_id, poa_anno,
                                                                       poas_subs_list, pag)

                    if poa:
                        notas_poa = GetNotasPoa(poa.id)
                    users_roots = getUsersRoots(estamento, user)

            if (poa and not poa.poa_include_subs) or len(poas_list) < 2:
                paginas = 1

    return [poa_redirect, estamentos_select, estamentos_list, estamentos_options_todos, can_edit, edit_notas, new_years,
            poa_years, poa, objetivos_list, poas_subs_list, cronogramas_list, notas_poa, tablero_list, poas_list,
            users_roots, paginas, paginado]


def GetPOAsLists(user, estamento, poa_anno, poa, poas, pag):
    current_user = user
    current_estamento = estamento
    poa_years = []
    for item_poa in poas:
        poa_years.append(item_poa.poa_anno)

    isCurrentUser = False
    isColaborador = False
    if user == estamento.estamento_user:
        isCurrentUser = True
    else:
        colaborador = Colaborador.objects.filter(colaborador_estamento=estamento, colaborador_user=user).first()
        if colaborador:
            isColaborador = True

    # if poa is None and user != estamento.estamento_user and IsRoot(estamento.id):
    if poa is None and not isCurrentUser and IsRoot(estamento.id):
        estamento_user = estamento.estamento_user
    else:
        estamento_user = user

    current_day = datetime.date.today()
    new_years = [current_day.year]
    if current_day.month >= MonthNewPoa():
        new_years.append(current_day.year + 1)

    estamentos_list = GetEstamentosList(current_user, estamento_user, False)
    if poa:
        if poa.poa_estamento.estamento_user != estamento_user:
            poa_estamentos_list = GetEstamentosList(current_user, poa.poa_estamento.estamento_user, False)
        else:
            poa_estamentos_list = estamentos_list
    else:
        poa_estamentos_list = estamentos_list

    estamentos_id_list = []
    estamentos = GetEstamentosList(current_user, estamento_user, False)
    for estamento in estamentos:
        if estamento.estamento_has_poa:
            estamentos_id_list.append(estamento.id)
        else:
            if IsRoot(estamento.id):
                estamentos_id_list.append(estamento.id)

    estamentos_subs = EstamentosSubsList(user, current_estamento.id, estamentos_list, poa_estamentos_list, poa_anno)
    estamentos_sub_list = estamentos_subs[0]
    estamentos_sub_list = list(estamentos_sub_list)
    poas_sub_list = estamentos_subs[1]
    poas_sub_list = list(poas_sub_list)
    tableros_sub_list = estamentos_subs[2]
    tableros_sub_list = list(tableros_sub_list)

    pag_total = 0
    paginado = 0
    if pag:
        paginado = int(Parametro.objects.filter(parametro_name='PAGINADO').first().parametro_value)
        pag_start = max((pag - 1) * paginado, 0)
        pag_end = pag * paginado
        if poa:
            pag_start = max(pag_start - 1, 0)
            pag_end = max(pag_end - 1, 0)
        pag_total = math.ceil(len(poas_sub_list) / paginado)
        poas_sub_list = poas_sub_list[pag_start:pag_end]

    estamentos_list = list(estamentos_list)
    currentUser = CustomUser.objects.get(id=user.id)
    if currentUser.is_superuser:
        estamentos_list += EstamentosAdminList()
    else:
        estamentos_list += estamentos_sub_list
    estamentos_list = EstamentosSortList(estamentos_list)

    estamentos_options_todos = ""
    count = len(estamentos_list)
    estamentos_select = []
    if count > 0:
        estamento = estamentos_list[0]
        estamento_tipo = estamento.estamento_tipo
        estamento_tipo_actual = estamento_tipo
        estamentos_options = str(estamento.id)
        estamentos_options_todos = estamentos_options
        estamento_tipo_id = estamento_tipo.id
        for index in range(1, count):
            estamento = estamentos_list[index]
            estamento_tipo = estamento.estamento_tipo
            estamentos_options_todos += "," + str(estamento.id)
            if estamento.estamento_tipo_id != estamento_tipo_id:
                estamento_tipo_id = estamento_tipo.id
                estamentos_select.append(
                    {"estamento_tipo": estamento_tipo_actual, "estamentos_options": estamentos_options})
                estamento_tipo_actual = estamento_tipo
                estamentos_options = str(estamento.id)
            else:
                estamentos_options += "," + str(estamento.id)

        if estamentos_options != "" and estamento_tipo_actual:
            estamentos_select.append(
                {"estamento_tipo": estamento_tipo_actual, "estamentos_options": estamentos_options})

    can_edit = False
    edit_notas = False
    if poa:
        if poa.poa_estado_id < 3 and poa.poa_anno >= current_day.year and poa.poa_estamento_id in estamentos_id_list:
            can_edit = True
        """
        if currentUser.is_superuser or (poa.poa_estado_id > 1 and (isCurrentUser or isColaborador)) or (
                poa.poa_estado_id <= 3 and poa.poa_estamento in estamentos_sub_list):
            edit_notas = True
        """
        if poa.poa_estado_id <= 3:
            edit_notas = True
    else:
        can_edit = IsPOAUser(estamento.id, user)

    return [estamentos_select, estamentos_list, poas_sub_list, tableros_sub_list, estamentos_options_todos, can_edit,
            edit_notas, new_years, poa_years, pag_total, paginado]


def CalcularPorcientos(actividad, presupuesto, porciento_meta, porciento_parcial_meta, peso_parcial_meta):
    presupuesto += actividad.actividad_presupuesto

    cronogramas = Cronograma.objects.filter(cronograma_actividad=actividad)
    cronogramas_terminados = Cronograma.objects.filter(cronograma_actividad=actividad,
                                                       cronograma_cumplimiento=True)
    presupuesto_gastado = 0
    for cronograma in cronogramas_terminados:
        presupuesto_gastado += cronograma.cronograma_presupuesto

    no_cronogramas = cronogramas.count()
    no_cronogramas_terminados = cronogramas_terminados.count()
    porciento = CalcularPorcentajes(actividad, no_cronogramas, no_cronogramas_terminados, porciento_meta)
    porciento_meta = porciento[0]
    porciento_actividad = porciento[1]

    current_month = datetime.date.today().month
    no_cronogramas_futuro = Cronograma.objects.filter(cronograma_actividad=actividad, cronograma_mes_id__gt=current_month).count()
    if no_cronogramas_terminados > 0:
        no_cronogramas_terminados = Cronograma.objects.filter(
            Q(cronograma_actividad=actividad, cronograma_cumplimiento=True) |
            Q(cronograma_actividad=actividad, cronograma_mes_id__gt=current_month)
        ).count()
    elif no_cronogramas_futuro == no_cronogramas:
        no_cronogramas_terminados = no_cronogramas_futuro

    porciento = CalcularPorcentajes(actividad, no_cronogramas, no_cronogramas_terminados, porciento_parcial_meta)
    porciento_parcial_meta = porciento[0]
    porciento_parcial_actividad = porciento[1]
    if no_cronogramas:
        peso_parcial_meta += actividad.actividad_peso

    return [presupuesto, presupuesto_gastado, porciento_meta, porciento_actividad, porciento_parcial_meta,
            porciento_parcial_actividad, peso_parcial_meta]


def CalcularPorcentajes(actividad, no_cronogramas, no_cronogramas_terminados, porciento_meta):
    porciento_actividad = 0
    if no_cronogramas > 0 and actividad.actividad_peso > 0:
        if no_cronogramas == no_cronogramas_terminados:
            porciento_actividad = 100
            porciento_peso = actividad.actividad_peso
        else:
            porciento_actividad = no_cronogramas_terminados * 100 / no_cronogramas
            porciento_actividad = round(porciento_actividad, 2)
            if porciento_actividad % 1 == 0:
                porciento_actividad = int(porciento_actividad)
            porciento_peso = porciento_actividad * actividad.actividad_peso / 100
            porciento_peso = round(porciento_peso, 2)
        porciento_meta += porciento_peso

    return [porciento_meta, porciento_actividad]


def ConvertHTMLtoDOC(request, template, poa_anno, poa, estamento, objetivos_list, is_current_user, doc_type):
    template = get_template(template)
    # host = r"35.185.0.227"  # static
    host = request.META['HTTP_HOST']  # dynamic
    ccs_bootstrap = r"http://" + host + r"/static/css/vendor/bootstrap.min.css"
    ccs_font = r"http://" + host + r"/static/css/vendor/font-awesome-pro.min.css"
    ccs_style = r"http://" + host + r"/static/css/style.css"
    styles = [ccs_bootstrap,
              ccs_font,
              ccs_style]

    htmlString = template.render({"styles": styles, "poa_anno": poa_anno, "poa": poa, "estamento": estamento,
                                  "objetivos_list": objetivos_list, "is_current_user": is_current_user,
                                  "edit_notas": False})

    if estamento:
        doc_name = RenameFile(estamento.estamento_name)
        doc_name = f'POA_{doc_name}_{str(poa_anno)}.{doc_type}'
    else:
        doc_name = f'POA_TMP.{doc_type}'

    doc_out = settings.MEDIA_ROOT + r"docs/" + doc_name

    if doc_type == 'docx':
        document = Document()
        sec = document.AddSection()
        paragraph = sec.AddParagraph()
        paragraph.AppendHTML(htmlString)
        section = document.Sections[0]
        section.PageSetup.Orientation = PageOrientation.Portrait
        document.SaveToFile(doc_out, FileFormat.Docx2016)
        document.Close()

        RemoveWarning(doc_out)

        doc_content = 'application/msword'
    else:
        # orientation = "portrait"
        orientation = "Landscape"
        options = {
            "encoding": "UTF-8",
            "orientation": orientation,
            "margin-top": "0.5in",
            "margin-left": "0in",
            "margin-bottom": "0.5in",
            "margin-right": "0in"
        }

        wkhtmltopdf_rute = Parametro.objects.filter(parametro_name='WKHTMLTOPDF_RUTE').first().parametro_value
        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_rute)

        pdfkit.from_string(htmlString, doc_out, options=options, configuration=config)

        doc_content = 'application/pdf'

    return {"doc_file": doc_out, "doc_content": doc_content}


def RemoveWarning(docx_file):
    from docx import Document
    doc = Document(docx_file)
    for paragraph in doc.paragraphs:
        if "Evaluation Warning: The document was created with Spire.Doc for Python." in paragraph.text:
            paragraph.clear()
    doc.save(docx_file)


def RenameFile(name):
    name = name.replace('á', 'a')
    name = name.replace('é', 'e')
    name = name.replace('í', 'i')
    name = name.replace('ó', 'o')
    name = name.replace('ú', 'u')
    name = name.replace('ñ', 'n')
    name = name.replace('Á', 'A')
    name = name.replace('É', 'E')
    name = name.replace('Í', 'I')
    name = name.replace('Ó', 'O')
    name = name.replace('Ú', 'U')
    name = name.replace('Ñ', 'N')
    name = re.sub(r'\W+', '_', name)

    return name


def MonthNewPoa():
    return int(Parametro.objects.filter(parametro_name='MES_INICIO').first().parametro_value)


def TestMail(request):
    template = "basic_email.html"
    user_name = "Moises"
    user_mail = "moisesmayet@uapa.edu.do"
    subject = "Su POA ha sido aprobado"
    message = "Felicidades, su POA(LABORATORIO DE INTELIGENCIA ARTIFICIAL Y BIGDATA) se ha aprobado, desde este momento no podrá realizar más cambios en el mismo "
    notes = "POA revisado "
    firma = GetURL(request) + "/static/images/logo/popin_firma.png"
    SendMail(user_name, user_mail, subject, template, message, notes, firma)
    return redirect("home")


def TestPDF(request):
    doc_out = ConvertHTMLtoDOC(request, "test_to_pdf.html", 2023, 0, None, [], True, 'pdf')

    return FileResponse(open(doc_out['doc_file'], 'rb'), content_type=doc_out['doc_content'])


def TestWord(request):
    doc_out = ConvertHTMLtoDOC(request, "test_to_pdf.html", 2023, 0, None, [], True, 'docx')

    return FileResponse(open(doc_out['doc_file'], 'rb'), content_type=doc_out['doc_content'])


def Test(request):
    object_list = ["Esto es una prueba"]
    return render(request, "test.html", {"object_list": object_list})


def IsRoot(estamento_root_id):
    child = Estamento.objects.filter(
        estamento_has_poa=True,
        estamento_roots__startswith=f'{estamento_root_id},',  # Comienza con 'id,'
    ).union(
        Estamento.objects.filter(
            estamento_has_poa=True,
            estamento_roots__contains=f',{estamento_root_id},'  # Contiene ',id,'
        )
    ).union(
        Estamento.objects.filter(
            estamento_has_poa=True,
            estamento_roots__endswith=f',{estamento_root_id}'  # Termina con ',id'
        )
    ).first()

    if child:
        return True
    return False


def IsChildOf(child_estamento, root_estamento_id):
    estamento_roots = list(map(int, child_estamento.estamento_roots.split(",")))
    if root_estamento_id in estamento_roots:
        return True
    return False


def GetColor(estamento, user):
    if estamento.estamento_user == user:
        return '#FF8300'
    elif Colaborador.objects.filter(colaborador_estamento=estamento, colaborador_user=user).first():
        return '#F7D164'

    return GetChildColor(estamento, user)


def GetChildColor(estamento, user):
    if Colaborador.objects.filter(colaborador_estamento=estamento,
                                  colaborador_user=user).first():
        return '#FFE066'
    elif estamento.id == 1 or Estamento.objects.filter(id=estamento.id, estamento_user=user).first():
        return '#49A1FF'

    return GetChildColor(estamento.estamento_sub, user)


def IncludeSubs(estamento_id, poa_anno):
    poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
    if poa:
        return poa.poa_include_subs
    else:
        return IsRoot(estamento_id)


def GetNotas(poa_id, nota_itemid, nota_itemname, nota_user):
    notas = Nota.objects.filter(nota_poa_id=poa_id, nota_itemid=nota_itemid, nota_itemname=nota_itemname,
                                nota_checked=False).annotate(
        nota_removeicon=Case(
            When(nota_user=nota_user, then=Value('fas fa-trash')),
            default=Value('fas fa-check'),
            output_field=CharField()
        ),
        nota_removetitle=Case(
            When(nota_user=nota_user, then=Value("Eliminar la nota")),
            default=Value("Marcar como realizada"),
            output_field=CharField()
        )
    ).order_by('nota_itemname', 'nota_date', 'id')
    return notas


def GetNotasPoa(poa_id):
    notas = Nota.objects.filter(nota_poa_id=poa_id).exclude(nota_itemname='evidencia').order_by('nota_date', 'id')
    return notas


def GetNotaValue(nota, count_notas, user, action, request_from):
    color = ""
    if action == "cheked":
        if count_notas == 0:
            if nota.nota_itemname == "objetivo" and request_from == 'poa_preview':
                color = "#F0F2F5"
            else:
                color = "#FFFFFF"
    elif action == "save":
        color = "rgba(246, 187, 186, 0.5)"

    if nota.nota_itemid:
        popover_id = f'popover_{nota.nota_itemname}_{nota.nota_itemid}'
    else:
        popover_id = f'popover_poa_{nota.nota_poa_id}'
    popover_nota_id = f'{popover_id}_{nota.id}'

    if nota.nota_user != user:
        nota_remove = False
        nota_removeicon = 'fas fa-check'
        nota_removetitle = "Marcar como realizada"
    else:
        nota_remove = True
        nota_removeicon = 'fas fa-trash'
        nota_removetitle = "Eliminar la nota"

    if nota.nota_itemname == 'objetivo':
        objetivo = ObjetivoOperativo.objects.filter(id=nota.nota_itemid).first()
        if objetivo:
            nota_order = f'[objetivo {objetivo.operativo_order}] '
    elif nota.nota_itemname == 'meta':
        meta = Meta.objects.filter(id=nota.nota_itemid).first()
        if meta:
            nota_order = f'[meta {meta.meta_operativo.operativo_order}.{meta.meta_order}] '
    elif nota.nota_itemname == 'actividad':
        actividad = Actividad.objects.filter(id=nota.nota_itemid).first()
        if actividad:
            nota_order = f'[actividad {actividad.actividad_meta.meta_operativo.operativo_order}.{actividad.actividad_meta.meta_order}.{actividad.actividad_order}] '

    return {'action': action, 'color': color, 'popover_id': popover_id, 'popover_nota_id': popover_nota_id,
            'nota_id': nota.id, 'nota_poa': nota.nota_poa_id,
            'nota_remove': nota_remove, 'nota_removeicon': nota_removeicon, 'nota_removetitle': nota_removetitle,
            'nota_user': nota.nota_user.username.title(), 'nota_date': nota.nota_date.strftime('%d %b %Y %H:%M'),
            'nota_description': nota.nota_description, 'nota_order': nota_order}


def UpdateNota(request, nota_description, nota_id, request_from):
    delete_nota = False
    count_notas = 0
    if nota_id:
        action = "cheked"
        nota = Nota.objects.filter(id=nota_id).first()
        if nota.nota_user != request.user:
            nota.nota_checked = True
            nota.save()
        else:
            delete_nota = True
            count_notas = -1
    else:
        action = "save"
        nota = CrearNota(request, nota_description, False, True)

    count_notas += Nota.objects.filter(nota_poa_id=nota.nota_poa_id, nota_itemid=nota.nota_itemid,
                                       nota_itemname=nota.nota_itemname,
                                       nota_checked=False).count()

    updatedValue = GetNotaValue(nota, count_notas, request.user, action, request_from)
    if delete_nota:
        nota.delete()

    return updatedValue


def CrearNota(request, nota_description, send_mail, send_notification):
    nota = Nota()
    nota.nota_user = request.user
    nota.nota_description = nota_description
    nota.nota_itemid = request.POST["nota_itemid"]
    nota.nota_itemname = request.POST["nota_itemname"]
    if "nota_itempoaid" in request.POST:
        nota.nota_poa_id = request.POST["nota_itempoaid"]
    else:
        nota.nota_poa_id = request.POST["poa_id"]

    if request.POST.get("nota_checked"):
        nota.nota_checked = True
    nota.save()

    if request.POST.get("nota_send"):
        try:
            poa = POA.objects.filter(id=nota.nota_poa_id).first()
            template = "basic_email.html"
            user_name = poa.poa_estamento.estamento_user.username
            user_mail = poa.poa_estamento.estamento_user.email
            subject = "Tienes una nueva notificación"
            message = nota.nota_description
            firma = GetURL(request) + "/static/images/logo/popin_firma.png"
            if send_mail:
                SendMail(user_name, user_mail, subject, template, message, "", firma)
            if send_notification:
                SendNotification(request.user, poa.poa_estamento.estamento_user, "ntf", message)
        except Exception as e:
            pass

    return nota


def RegistrarLog(estamento_id, poa_anno, log_user, log_action, log_description):
    log = Log()
    if estamento_id and poa_anno:
        poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
        log.log_poa = poa.id
        log.log_poaname = poa.poa_estamento.estamento_name
    log.log_user = log_user.id
    log.log_username = log_user.username
    log.log_action = log_action
    log.log_description = log_description
    log.save()


def GetEvidenciaType(evidencia_id):
    evidencia = Evidencia.objects.get(id=evidencia_id)
    evidencia_type = ''
    evidencia_icon = ''
    evidencia_preview = False

    if evidencia:
        extension = os.path.splitext(evidencia.evidencia_file.name)[1]
        if extension.lower() in '.jpg,.jpeg,.gif,.png,.tiff,.tif,.RAW,.bmp,.psd,.eps,.pic':
            evidencia_type = 'image'
            evidencia_icon = '/static/images/icons/imagen.png'
            evidencia_preview = True
        else:
            if extension.lower() in '.avi,.wmv,.asf,.mov,.flv,.rm,.rmvb,.mp4,.mkv,.mks,.3gpp,.mpg':
                evidencia_type = 'video'
                evidencia_icon = '/static/images/icons/video.png'
                evidencia_preview = True
            elif extension.lower() in '.mp3,.wav,.cda,.midi,.mp3.ogg,.wma':
                evidencia_type = 'audio'
                evidencia_icon = '/static/images/icons/audio.png'
            elif extension.lower() in '.doc,.docx,.docm,.dotx,.dotm,.odt ':
                evidencia_type = 'office'
                evidencia_icon = '/static/images/icons/word.png'
                evidencia_preview = True
            elif extension.lower() in '.xls,.xlsx,.xlsm,.xltx,.xltm,.xlsb,.xlam,.ods ':
                evidencia_type = 'office'
                evidencia_icon = '/static/images/icons/excel.png'
                evidencia_preview = True
            elif extension.lower() in '.ppt,.pptx,.pptm,.potx,.potm,.ppam,.ppsx,.ppsm,.sldx,.sldm,.thmx,.odp':
                evidencia_type = 'office'
                evidencia_icon = '/static/images/icons/ppt.png'
                evidencia_preview = True
            elif extension.lower() in '.zip,.gz,.gzip,.rar,.tar,.tgz,.zip,.iso':
                evidencia_type = 'rar'
                evidencia_icon = '/static/images/icons/rar.png'
            elif extension.lower() == '.pdf':
                evidencia_type = 'pdf'
                evidencia_icon = '/static/images/icons/pdf.png'
                evidencia_preview = True
            elif extension.lower() == '.txt':
                evidencia_type = 'texto'
                evidencia_icon = '/static/images/icons/texto.png'
            else:
                evidencia_type = 'otro'
                evidencia_icon = '/static/images/icons/otro.png'

    return {'evidencia_type': evidencia_type, 'evidencia_icon': evidencia_icon, 'evidencia_preview': evidencia_preview}


def GetEstamentosList(user, estamento_user, check_has_poa):
    colaboradores_ids = Colaborador.objects.filter(colaborador_user=user).values_list(
        'colaborador_estamento', flat=True
    )
    if user != estamento_user:
        estamento_main = False
    else:
        estamento_main = True

    if check_has_poa:
        estamentos_main = Estamento.objects.filter(estamento_user=estamento_user, estamento_has_poa=True).exclude(
            id__in=colaboradores_ids).annotate(estamento_main=Value(estamento_main, output_field=BooleanField()))
    else:
        estamentos_main = Estamento.objects.filter(estamento_user=estamento_user).exclude(
            id__in=colaboradores_ids).annotate(estamento_main=Value(estamento_main, output_field=BooleanField()))
    estamentos_colaboradores = Estamento.objects.filter(id__in=colaboradores_ids).annotate(
        estamento_main=Value(False, output_field=BooleanField()))
    estamentos = list(estamentos_main.union(estamentos_colaboradores, all=True))

    return estamentos


def getUsersRoots(estamento, user):
    colaboradores = Colaborador.objects.filter(colaborador_estamento=estamento)
    users_root = getUsersColaboradores(colaboradores)
    if estamento.estamento_user != user:
        users_root += [estamento.estamento_user]
    else:
        estamento_roots = list(map(int, estamento.estamento_roots.split(","))) if estamento.estamento_roots else []
        estamento_roots = [est_id for est_id in estamento_roots if est_id != 1]
        for estamento_id in estamento_roots:
            root = Estamento.objects.filter(id=estamento_id).first()
            if root:
                users_root.append(root.estamento_user)
                colaboradores = Colaborador.objects.filter(colaborador_estamento=root)
                users_root += getUsersColaboradores(colaboradores)

    return users_root


def GetPOAUsers(poa_id):
    poa = POA.objects.filter(id=poa_id).first()
    users_list = []
    if poa:
        users_list.append(poa.poa_estamento.estamento_user)
        colaboradores = Colaborador.objects.filter(colaborador_estamento_id=poa.poa_estamento_id)
        users_list += [colaborador.colaborador_user for colaborador in colaboradores]

    return users_list


def IsPOAUser(estamento_id, user):
    estamento = Estamento.objects.filter(id=estamento_id, estamento_user=user).first()
    colaborador = Colaborador.objects.filter(colaborador_estamento_id=estamento_id, colaborador_user=user).first()
    if estamento or colaborador:
        return True
    return False


def getUsersColaboradores(colaboradores):
    users = []
    for colaborador in colaboradores:
        users.append(colaborador.colaborador_user)
    return users


def GetPOAList(request, user, estamento_id, poa_anno, pag, key):
    poas = GetPOAs(user, estamento_id, poa_anno, pag, key)
    return poas


def GetFormatPesos(amount):
    try:
        amount_float = float(amount)
        return "RD$ {:,.0f}".format(amount_float)
    except ValueError:
        return "RD$ 0"


def GetFormatPorciento(porciento):
    if porciento % 1 == 0:
        porciento = int(porciento)
    return porciento


def OrderAll(estamento_id, poa_anno, user, request_from):
    poa = POA.objects.filter(poa_estamento_id=estamento_id, poa_anno=poa_anno).first()
    objetivos = ObjetivoOperativo.objects.filter(operativo_poa=poa)
    operativo_order = 1
    for objetivo in objetivos:
        objetivo.operativo_order = operativo_order
        objetivo.save()
        operativo_order += 1
        metas = Meta.objects.filter(meta_operativo=objetivo)
        meta_order = 1
        for meta in metas:
            meta.meta_order = meta_order
            meta.save()
            meta_order += 1
            actividad_order = 1
            actividades = Actividad.objects.filter(actividad_meta=meta)
            for actividad in actividades:
                actividad.actividad_order = actividad_order
                actividad.save()
                actividad_order += 1

    RegistrarLog(estamento_id, poa_anno, user, "Order",
                 f"Se reorganizaron todos los índices desde {request_from}")
